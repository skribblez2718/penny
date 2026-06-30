"""
title: Word Document Generator (Standalone)
author: sketch (converted from Open WebUI Pipe)
version: 2.0.0
license: MIT
description: Programmatic Word document (.docx) generation from structured content.
             Supports lightweight markdown-to-docx (DocumentContent) and full-control
             typed API (WordInstructions). No Open WebUI dependency.
requirements: python-docx,pydantic,lxml
"""

from __future__ import annotations

import base64
import io
import json
import os
import re
import uuid
from datetime import datetime
from typing import Any, Dict, List, Literal, Optional, Tuple, Union

from pydantic import BaseModel, Field, field_validator, model_validator

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# Pydantic Schemas
# ============================================================


# ---------- Shared ----------


class ImageSpec(BaseModel):
    """Image to insert, by filename or bytes via base64."""

    name: Optional[str] = Field(
        default=None, description="File name as uploaded/attached."
    )
    b64: Optional[str] = Field(
        default=None, description="Base64-encoded binary of the image."
    )
    width_inches: float = Field(default=2.0, ge=0.1, le=20.0)
    height_inches: Optional[float] = Field(default=None, ge=0.1, le=20.0)


# ---------- Word (DOCX) ----------


class RunSpec(BaseModel):
    """A single text run with its own formatting within a paragraph."""

    text: str = Field(..., description="Run text content.")
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = Field(default=None, ge=4.0, le=120.0)
    font_color: Optional[str] = Field(
        default=None, description="Hex color WITHOUT #, e.g. '2F5496'."
    )
    highlight_color: Optional[str] = Field(
        default=None, description="Highlight color name e.g. 'yellow', 'cyan'."
    )
    superscript: Optional[bool] = None
    subscript: Optional[bool] = None
    strike: Optional[bool] = None


class ParagraphSpec(BaseModel):
    """Paragraph with full formatting support."""

    text: str = Field(
        default="",
        description="Simple text (backward compat). Ignored if 'runs' is provided.",
    )
    runs: List[RunSpec] = Field(
        default_factory=list,
        description="Rich runs for mixed formatting in one paragraph.",
    )
    bold: bool = False
    italic: bool = False
    underline: bool = False
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = Field(default=None, ge=4.0, le=120.0)
    font_color: Optional[str] = Field(
        default=None, description="Hex color e.g. '2F5496'."
    )
    style: Optional[str] = Field(
        default=None,
        description="Word paragraph style e.g. 'Heading 1', 'List Bullet'.",
    )
    alignment: Optional[str] = Field(
        default=None, description="'left', 'center', 'right', or 'justify'."
    )
    space_before_pt: Optional[float] = Field(default=None, ge=0, le=100)
    space_after_pt: Optional[float] = Field(default=None, ge=0, le=100)
    line_spacing: Optional[float] = Field(
        default=None, ge=0.5, le=5.0, description="Line spacing multiplier."
    )
    line_spacing_twips: Optional[int] = Field(
        default=None, description="Exact line spacing in twips. Overrides line_spacing."
    )
    indent_left_inches: Optional[float] = Field(default=None, ge=0, le=6)
    indent_right_inches: Optional[float] = Field(default=None, ge=0, le=6)
    indent_first_line_inches: Optional[float] = Field(default=None, ge=-3, le=6)
    page_break_before: bool = False
    keep_with_next: bool = False
    keep_together: bool = False
    outline_level: Optional[int] = Field(default=None, ge=0, le=8)


class TableCellSpec(BaseModel):
    """Rich cell specification for advanced tables."""

    text: str = ""
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    font_size_pt: Optional[float] = None
    font_color: Optional[str] = None
    font_name: Optional[str] = None
    alignment: Optional[str] = None
    vertical_alignment: Optional[str] = Field(
        default=None, description="'top','center','bottom'"
    )
    shading_color: Optional[str] = Field(
        default=None, description="Hex fill color e.g. 'D9D9D9'."
    )
    merge_right: int = Field(
        default=0, ge=0, description="Number of cells to merge rightward."
    )
    merge_down: int = Field(
        default=0, ge=0, description="Number of cells to merge downward."
    )
    runs: List[RunSpec] = Field(
        default_factory=list, description="Rich runs for mixed formatting."
    )


class TableSpec(BaseModel):
    """Table with full formatting support."""

    rows: List[List[str]] = Field(default_factory=list, description="2D string data.")
    rich_rows: List[List[TableCellSpec]] = Field(
        default_factory=list,
        description="Fully formatted rows. Takes precedence over 'rows'.",
    )
    style: Optional[str] = Field(default=None, description="Word table style name.")
    header_shading: Optional[str] = Field(
        default=None, description="Hex fill for first row header e.g. 'D9D9D9'."
    )
    header_bold: bool = Field(default=True, description="Bold first row.")

    @field_validator("rows", mode="before")
    @classmethod
    def filter_none_cells(cls, v: Any) -> list:
        if isinstance(v, list):
            return [
                [cell if cell is not None else "" for cell in row]
                for row in v
                if isinstance(row, list)
            ]
        return v

    column_widths_inches: List[float] = Field(
        default_factory=list, description="Column widths in inches."
    )
    borders: bool = Field(default=True, description="Enable cell borders.")
    border_color: str = Field(default="999999", description="Hex border color.")
    border_size_pt: float = Field(
        default=0.5, description="Border thickness in points."
    )
    table_alignment: Optional[str] = Field(
        default=None, description="'left','center','right'."
    )


class FindReplaceSpec(BaseModel):
    """Naive find/replace across paragraph runs."""

    find: str
    replace: str


class HeaderFooterContentSpec(BaseModel):
    """Content specification for a header or footer."""

    text: Optional[str] = Field(default=None, description="Simple text.")
    runs: List[RunSpec] = Field(
        default_factory=list, description="Rich formatted runs."
    )
    alignment: Optional[str] = Field(
        default=None, description="'left','center','right'."
    )
    font_size_pt: Optional[float] = None
    font_name: Optional[str] = None
    include_page_number: bool = Field(default=False, description="Append page number.")
    images: List[ImageSpec] = Field(
        default_factory=list, description="Images to insert into this header/footer."
    )


class SectionSpec(BaseModel):
    """Document section with its own page layout, margins, and headers/footers."""

    page_width_inches: float = Field(default=8.5)
    page_height_inches: float = Field(default=11.0)
    orientation: Optional[str] = Field(
        default=None, description="'portrait' or 'landscape'."
    )
    top_margin_inches: Optional[float] = Field(default=None, ge=0, le=5)
    bottom_margin_inches: Optional[float] = Field(default=None, ge=0, le=5)
    left_margin_inches: Optional[float] = Field(default=None, ge=0, le=5)
    right_margin_inches: Optional[float] = Field(default=None, ge=0, le=5)
    header_distance_inches: Optional[float] = Field(default=None, ge=0, le=3)
    footer_distance_inches: Optional[float] = Field(default=None, ge=0, le=3)
    different_first_page: bool = Field(
        default=False, description="Enable different first-page header/footer."
    )
    header: Optional[HeaderFooterContentSpec] = None
    footer: Optional[HeaderFooterContentSpec] = None
    first_page_header: Optional[HeaderFooterContentSpec] = None
    first_page_footer: Optional[HeaderFooterContentSpec] = None
    even_page_header: Optional[HeaderFooterContentSpec] = None
    even_page_footer: Optional[HeaderFooterContentSpec] = None


class TocSpec(BaseModel):
    """Table of Contents specification."""

    title: str = Field(default="Table of Contents")
    title_style: Optional[str] = Field(
        default="TOC Heading", description="Style for TOC heading."
    )
    max_level: int = Field(
        default=2, ge=1, le=9, description="Max heading level to include."
    )
    tab_position_twips: int = Field(
        default=9360, description="Tab stop for page numbers."
    )


class ContentBlock(BaseModel):
    """A single content element in the document."""

    type: Literal[
        "paragraph", "table", "image", "page_break", "section_break", "toc"
    ] = Field(..., description="Block type.")
    paragraph: Optional[ParagraphSpec] = None
    table: Optional[TableSpec] = None
    image: Optional[ImageSpec] = None
    section: Optional[SectionSpec] = None
    toc: Optional[TocSpec] = None


class WordStyleOverride(BaseModel):
    """Override a built-in Word style."""

    style_id: str = Field(..., description="e.g. 'Heading1', 'Normal'.")
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    font_color: Optional[str] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    line_spacing: Optional[float] = None
    line_spacing_twips: Optional[int] = None
    alignment: Optional[str] = None
    outline_level: Optional[int] = None


class WordInstructions(BaseModel):
    """Full Word document instructions."""

    content: List[ContentBlock] = Field(
        default_factory=list, description="Ordered content blocks."
    )
    paragraphs: List[ParagraphSpec] = Field(default_factory=list)
    tables: List[TableSpec] = Field(default_factory=list)
    images: List[ImageSpec] = Field(default_factory=list)
    sections: List[SectionSpec] = Field(
        default_factory=list, description="Section properties."
    )
    style_overrides: List[WordStyleOverride] = Field(
        default_factory=list, description="Override built-in styles."
    )
    default_font: Optional[str] = Field(
        default=None, description="Default document font."
    )
    default_font_size_pt: Optional[float] = Field(
        default=None, description="Default font size."
    )
    header_text: Optional[str] = None
    footer_text: Optional[str] = None
    find_replace: List[FindReplaceSpec] = Field(default_factory=list)
    default_style: Optional[str] = Field(default=None)


# ---------- General document content (lightweight markdown-to-docx) ----------


class DocumentSection(BaseModel):
    """A single section with a heading and markdown-style content."""

    heading: str = Field(
        default="", description="Section heading (rendered as Heading 1)."
    )
    content: str = Field(
        default="",
        description=(
            "Section body as lightweight markdown. Supports: ## H2, ### H3, #### H4, "
            "- bullets, 1. numbered lists, **bold**, *italic*, | table rows |, "
            "and plain paragraphs separated by blank lines."
        ),
    )

    @field_validator("heading", "content", mode="before")
    @classmethod
    def coerce_none_to_empty(cls, v: Any) -> str:
        return v if v is not None else ""


class DocumentContent(BaseModel):
    """
    Lightweight general document specification. Pass content as
    markdown-style text and optional style hints; the tool builds the full
    .docx internally.
    """

    title: str = Field(
        default="", description="Document title (rendered as Title style)."
    )
    subtitle: Optional[str] = Field(default=None, description="Optional subtitle.")
    sections: List[DocumentSection] = Field(
        default_factory=list,
        description="Ordered list of document sections with headings and markdown content.",
    )

    @field_validator("sections", mode="before")
    @classmethod
    def filter_none_sections(cls, v: Any) -> list:
        if isinstance(v, list):
            return [item for item in v if item is not None]
        return v

    @model_validator(mode="before")
    @classmethod
    def parse_markdown_field(cls, data: Any) -> Any:
        """If 'markdown' is provided as raw text, parse it into title, subtitle, and sections."""
        if not isinstance(data, dict) or "markdown" not in data:
            return data

        markdown_text = data.pop("markdown")
        if not markdown_text or not markdown_text.strip():
            return data

        lines = markdown_text.split("\n")

        # Extract H1 as title
        title = ""
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()
            if stripped.startswith("# ") and not stripped.startswith("## "):
                title = stripped[2:]
                i += 1
                break
            i += 1

        # Collect everything between title and first H2 as subtitle
        subtitle_start = i
        while i < len(lines):
            stripped = lines[i].strip()
            if stripped.startswith("## "):
                break
            i += 1
        subtitle_text = "\n".join(lines[subtitle_start:i]).strip()

        # Parse H2-delimited sections
        sections: list = []
        current_heading = ""
        current_content_lines: List[str] = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if stripped.startswith("## ") and not stripped.startswith("### "):
                if current_heading:
                    sections.append(DocumentSection(
                        heading=current_heading,
                        content="\n".join(current_content_lines).strip()
                    ))
                current_heading = stripped[3:]
                current_content_lines = []
            elif current_heading:
                current_content_lines.append(line)
            i += 1

        # Final section
        if current_heading:
            sections.append(DocumentSection(
                heading=current_heading,
                content="\n".join(current_content_lines).strip()
            ))

        # Only fill in if not explicitly provided
        if "title" not in data or not data.get("title"):
            data["title"] = title
        if "subtitle" not in data or data.get("subtitle") is None:
            data["subtitle"] = subtitle_text if subtitle_text else None
        if "sections" not in data or not data.get("sections"):
            data["sections"] = sections

        return data


# ============================================================
# Resume-Specific Models
# ============================================================


class ResumeContent(BaseModel):
    """
    Resume-specific content specification. Reads a markdown resume
    and generates a .docx matching the canonical resume template layout.
    
    The markdown is parsed semantically — section headings determine
    element types (skills labels → bold prefixes, role titles → special
    formatting, etc.) for exact layout reproduction.
    """

    markdown: str = Field(
        default="",
        description="Full markdown resume content to convert to .docx."
    )
    font_family: str = Field(
        default="Arial",
        description="Base font family for all body text."
    )
    font_size_pt: float = Field(
        default=9.0,
        description="Base font size in points for body text."
    )
    header_color: str = Field(
        default="245291",
        description="Hex color for section headers (without #)."
    )
    name_color: str = Field(
        default="245291",
        description="Hex color for the name header (without #)."
    )
    border_color: str = Field(
        default="D9E2F3",
        description="Hex color for section header bottom border (without #)."
    )

    # --- Style overrides (all optional — sensible defaults applied) ---
    font: Optional[str] = Field(
        default=None,
        description="Default body font name, e.g. 'Calibri', 'Times New Roman', 'Arial'.",
    )
    font_size_pt: Optional[float] = Field(
        default=None, ge=6.0, le=72.0, description="Default body font size in points."
    )
    heading_font: Optional[str] = Field(
        default=None, description="Font for headings. Defaults to body font if omitted."
    )
    heading_color: Optional[str] = Field(
        default=None,
        description="Hex color for headings WITHOUT #, e.g. '2F5496'. Default: '2F5496'.",
    )
    line_spacing: Optional[float] = Field(
        default=None,
        ge=0.5,
        le=3.0,
        description="Line spacing multiplier. Default: 1.15.",
    )
    orientation: Optional[str] = Field(
        default=None, description="'portrait' (default) or 'landscape'."
    )
    margin_inches: Optional[float] = Field(
        default=None,
        ge=0.3,
        le=3.0,
        description="Uniform page margin in inches. Default: 1.0.",
    )
    header_text: Optional[str] = Field(
        default=None, description="Text for page header."
    )
    footer_text: Optional[str] = Field(
        default=None, description="Text for page footer."
    )
    include_page_numbers: bool = Field(
        default=True, description="Show page numbers in footer."
    )
    include_toc: bool = Field(
        default=False, description="Insert a Table of Contents after the title."
    )
    table_style: Optional[str] = Field(
        default=None,
        description="Table border/shading preset: 'grid' (default), 'minimal', 'none'.",
    )


# ============================================================
# Utilities
# ============================================================


def _now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _choose_output_name(base: Optional[str]) -> str:
    root = base or f"document_{_now_stamp()}"
    return f"{root}.docx"


def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    if len(h) == 6:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return RGBColor(0, 0, 0)


def _alignment_enum(align_str: Optional[str]):
    if not align_str:
        return None
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align_str.lower())


# ============================================================
# Word (.docx) Core Builders
# ============================================================


def _apply_run_formatting(run, spec: RunSpec) -> None:
    if spec.bold is not None:
        run.bold = spec.bold
    if spec.italic is not None:
        run.italic = spec.italic
    if spec.underline is not None:
        run.underline = spec.underline
    if spec.font_name:
        run.font.name = spec.font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), spec.font_name)
    if spec.font_size_pt is not None:
        run.font.size = Pt(spec.font_size_pt)
    if spec.font_color:
        run.font.color.rgb = _hex_to_rgb(spec.font_color)
    if spec.highlight_color:
        try:
            from docx.enum.text import WD_COLOR_INDEX

            color_map = {
                "yellow": WD_COLOR_INDEX.YELLOW,
                "green": WD_COLOR_INDEX.BRIGHT_GREEN,
                "cyan": WD_COLOR_INDEX.TURQUOISE,
                "pink": WD_COLOR_INDEX.PINK,
                "red": WD_COLOR_INDEX.RED,
                "blue": WD_COLOR_INDEX.BLUE,
                "gray": WD_COLOR_INDEX.GRAY_25,
            }
            run.font.highlight_color = color_map.get(spec.highlight_color.lower())
        except Exception:
            pass
    if spec.superscript:
        run.font.superscript = True
    if spec.subscript:
        run.font.subscript = True
    if spec.strike:
        run.font.strike = True


def _apply_paragraph_formatting(para, spec: ParagraphSpec) -> None:
    if spec.style:
        try:
            para.style = spec.style
        except Exception:
            pass
    if spec.alignment:
        align = _alignment_enum(spec.alignment)
        if align is not None:
            para.alignment = align
    pf = para.paragraph_format
    if spec.space_before_pt is not None:
        pf.space_before = Pt(spec.space_before_pt)
    if spec.space_after_pt is not None:
        pf.space_after = Pt(spec.space_after_pt)
    if spec.line_spacing_twips is not None:
        pPr = para._element.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(spacing)
        spacing.set(qn("w:line"), str(spec.line_spacing_twips))
        spacing.set(qn("w:lineRule"), "exact")
    elif spec.line_spacing is not None:
        pf.line_spacing = spec.line_spacing
    if spec.indent_left_inches is not None:
        pf.left_indent = Inches(spec.indent_left_inches)
    if spec.indent_right_inches is not None:
        pf.right_indent = Inches(spec.indent_right_inches)
    if spec.indent_first_line_inches is not None:
        if spec.indent_first_line_inches >= 0:
            pf.first_line_indent = Inches(spec.indent_first_line_inches)
        else:
            pPr = para._element.get_or_add_pPr()
            ind = pPr.find(qn("w:ind"))
            if ind is None:
                ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
                pPr.append(ind)
            ind.set(
                qn("w:hanging"), str(int(abs(spec.indent_first_line_inches) * 1440))
            )
    if spec.page_break_before:
        pf.page_break_before = True
    if spec.keep_with_next:
        pf.keep_with_next = True
    if spec.keep_together:
        pf.keep_together = True
    if spec.outline_level is not None:
        pPr = para._element.get_or_add_pPr()
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = parse_xml(
                f'<w:outlineLvl {nsdecls("w")} w:val="{spec.outline_level}"/>'
            )
            pPr.append(outline)
        else:
            outline.set(qn("w:val"), str(spec.outline_level))


def _add_paragraph(doc, spec: ParagraphSpec):
    para = doc.add_paragraph()
    _apply_paragraph_formatting(para, spec)
    if spec.runs:
        for rs in spec.runs:
            run = para.add_run(rs.text)
            _apply_run_formatting(run, rs)
    else:
        run = para.add_run(spec.text)
        simple_run = RunSpec(
            text=spec.text,
            bold=spec.bold or None,
            italic=spec.italic or None,
            underline=spec.underline or None,
            font_name=spec.font_name,
            font_size_pt=spec.font_size_pt,
            font_color=spec.font_color,
        )
        _apply_run_formatting(run, simple_run)
    return para


def _set_cell_shading(cell, color_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shading)


def _set_cell_borders(cell, color: str = "999999", size: int = 4) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders_xml = f"""<w:tcBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
        <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
        <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
    </w:tcBorders>"""
    tcPr.append(parse_xml(borders_xml))


def _set_cell_width(cell, width_twips: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
        tcPr.append(tcW)
    else:
        tcW.set(qn("w:w"), str(width_twips))
        tcW.set(qn("w:type"), "dxa")


def _format_cell_text(
    cell,
    text,
    bold=None,
    italic=None,
    font_size_pt=None,
    font_color=None,
    font_name=None,
    alignment=None,
    runs=None,
):
    if runs:
        para = cell.paragraphs[0]
        para.clear()
        for rs in runs:
            run = para.add_run(rs.text)
            _apply_run_formatting(run, rs)
    else:
        cell.text = str(text)
        if cell.paragraphs and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if font_size_pt is not None:
                run.font.size = Pt(font_size_pt)
            if font_color:
                run.font.color.rgb = _hex_to_rgb(font_color)
            if font_name:
                run.font.name = font_name
    if alignment and cell.paragraphs:
        align = _alignment_enum(alignment)
        if align is not None:
            cell.paragraphs[0].alignment = align


def _add_table(doc, spec: TableSpec):
    use_rich = bool(spec.rich_rows)
    if use_rich:
        rows_data = spec.rich_rows
        num_rows = len(rows_data)
        num_cols = max(len(r) for r in rows_data) if rows_data else 0
    else:
        rows_data = spec.rows
        num_rows = len(rows_data)
        num_cols = max(len(r) for r in rows_data) if rows_data else 0

    if num_rows == 0 or num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)

    if spec.style:
        try:
            table.style = spec.style
        except Exception:
            pass

    if spec.table_alignment:
        tbl_align_map = {
            "left": WD_TABLE_ALIGNMENT.LEFT,
            "center": WD_TABLE_ALIGNMENT.CENTER,
            "right": WD_TABLE_ALIGNMENT.RIGHT,
        }
        try:
            table.alignment = tbl_align_map.get(spec.table_alignment.lower())
        except Exception:
            pass

    col_widths_twips = []
    if spec.column_widths_inches:
        col_widths_twips = [int(w * 1440) for w in spec.column_widths_inches]

    border_color = spec.border_color.lstrip("#") if spec.border_color else "999999"
    border_size = int(spec.border_size_pt * 8)

    for i in range(num_rows):
        for j in range(num_cols):
            cell = table.cell(i, j)
            if use_rich:
                if j < len(rows_data[i]):
                    cs = rows_data[i][j]
                    _format_cell_text(
                        cell,
                        cs.text,
                        bold=cs.bold,
                        italic=cs.italic,
                        font_size_pt=cs.font_size_pt,
                        font_color=cs.font_color,
                        font_name=cs.font_name,
                        alignment=cs.alignment,
                        runs=cs.runs if cs.runs else None,
                    )
                    if cs.shading_color:
                        _set_cell_shading(cell, cs.shading_color.lstrip("#"))
                    if cs.vertical_alignment:
                        va_map = {"top": "top", "center": "center", "bottom": "bottom"}
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        vAlign = parse_xml(
                            f'<w:vAlign {nsdecls("w")} w:val="{va_map.get(cs.vertical_alignment.lower(), "top")}"/>'
                        )
                        tcPr.append(vAlign)
            else:
                if j < len(rows_data[i]):
                    cell.text = str(rows_data[i][j])
            if i == 0:
                if spec.header_shading:
                    _set_cell_shading(cell, spec.header_shading.lstrip("#"))
                if spec.header_bold and cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            if col_widths_twips and j < len(col_widths_twips):
                _set_cell_width(cell, col_widths_twips[j])
            if spec.borders:
                _set_cell_borders(cell, border_color, border_size)

    return table


def _add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br)
    return para


def _add_section_break(doc, spec: Optional[SectionSpec] = None):
    section = doc.add_section()
    if spec:
        _apply_section_props(section, spec)
    return section


def _apply_section_props(section, spec: SectionSpec) -> None:
    if spec.page_width_inches:
        section.page_width = Inches(spec.page_width_inches)
    if spec.page_height_inches:
        section.page_height = Inches(spec.page_height_inches)
    if spec.orientation:
        if spec.orientation.lower() == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            if section.page_width < section.page_height:
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )
        else:
            section.orientation = WD_ORIENT.PORTRAIT
    if spec.top_margin_inches is not None:
        section.top_margin = Inches(spec.top_margin_inches)
    if spec.bottom_margin_inches is not None:
        section.bottom_margin = Inches(spec.bottom_margin_inches)
    if spec.left_margin_inches is not None:
        section.left_margin = Inches(spec.left_margin_inches)
    if spec.right_margin_inches is not None:
        section.right_margin = Inches(spec.right_margin_inches)
    if spec.header_distance_inches is not None:
        section.header_distance = Inches(spec.header_distance_inches)
    if spec.footer_distance_inches is not None:
        section.footer_distance = Inches(spec.footer_distance_inches)
    if spec.different_first_page:
        section.different_first_page_header_footer = True
    _apply_header_footer(section, spec)


def _apply_hf_content(hf_obj, spec: HeaderFooterContentSpec) -> None:
    for idx, img in enumerate(spec.images):
        if img.b64:
            image_bytes = base64.b64decode(img.b64)
            if idx == 0 and hf_obj.paragraphs:
                img_para = hf_obj.paragraphs[0]
            else:
                img_para = hf_obj.add_paragraph()
            run = img_para.add_run()
            image_stream = io.BytesIO(image_bytes)
            width = Inches(img.width_inches)
            height = Inches(img.height_inches) if img.height_inches else None
            inline = run.part.new_pic_inline(image_stream, width, height)
            drawing = parse_xml(f'<w:drawing {nsdecls("w")}/>')
            drawing.append(inline)
            run._r.append(drawing)

    has_text = spec.text or spec.runs or spec.include_page_number
    if has_text or not spec.images:
        if spec.images:
            para = hf_obj.add_paragraph()
        else:
            para = hf_obj.paragraphs[0] if hf_obj.paragraphs else hf_obj.add_paragraph()
        if spec.alignment:
            align = _alignment_enum(spec.alignment)
            if align is not None:
                para.alignment = align
        if spec.runs:
            for rs in spec.runs:
                run = para.add_run(rs.text)
                _apply_run_formatting(run, rs)
        elif spec.text:
            run = para.add_run(spec.text)
            if spec.font_size_pt:
                run.font.size = Pt(spec.font_size_pt)
            if spec.font_name:
                run.font.name = spec.font_name
        if spec.include_page_number:
            run = para.add_run()
            fldChar_begin = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
            )
            run._element.append(fldChar_begin)
            run2 = para.add_run()
            instrText = parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
            )
            run2._element.append(instrText)
            run3 = para.add_run()
            fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run3._element.append(fldChar_end)


def _apply_header_footer(section, spec: SectionSpec) -> None:
    if spec.header:
        header = section.header
        header.is_linked_to_previous = False
        _apply_hf_content(header, spec.header)
    if spec.footer:
        footer = section.footer
        footer.is_linked_to_previous = False
        _apply_hf_content(footer, spec.footer)
    if spec.different_first_page:
        if spec.first_page_header:
            fp_header = section.first_page_header
            _apply_hf_content(fp_header, spec.first_page_header)
        if spec.first_page_footer:
            fp_footer = section.first_page_footer
            _apply_hf_content(fp_footer, spec.first_page_footer)
    if spec.even_page_header:
        ep_header = section.even_page_header
        _apply_hf_content(ep_header, spec.even_page_header)
    if spec.even_page_footer:
        ep_footer = section.even_page_footer
        _apply_hf_content(ep_footer, spec.even_page_footer)


def _add_toc(doc, spec: TocSpec):
    heading_para = doc.add_paragraph()
    if spec.title_style:
        try:
            heading_para.style = spec.title_style
        except Exception:
            try:
                heading_para.style = "Heading 1"
            except Exception:
                pass
    run = heading_para.add_run(spec.title)

    para = doc.add_paragraph()
    run1 = para.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1._element.append(fldChar_begin)
    run2 = para.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-{spec.max_level}" \\h \\z \\u </w:instrText>'
    )
    run2._element.append(instrText)
    run3 = para.add_run()
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_sep)
    run4 = para.add_run(
        "[Update Table of Contents: select all with Ctrl+A then press F9]"
    )
    run4.italic = True
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run5 = para.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar_end)
    return para


def _apply_style_overrides(doc, overrides: List[WordStyleOverride]) -> None:
    for ovr in overrides:
        try:
            style_name_map = {
                "Heading1": "Heading 1",
                "Heading2": "Heading 2",
                "Heading3": "Heading 3",
                "Normal": "Normal",
                "Title": "Title",
                "Subtitle": "Subtitle",
                "ListBullet": "List Bullet",
                "ListNumber": "List Number",
                "TOCHeading": "TOC Heading",
            }
            style_name = style_name_map.get(ovr.style_id, ovr.style_id)
            style = doc.styles[style_name]
            if ovr.font_name:
                style.font.name = ovr.font_name
            if ovr.font_size_pt is not None:
                style.font.size = Pt(ovr.font_size_pt)
            if ovr.font_color:
                style.font.color.rgb = _hex_to_rgb(ovr.font_color)
            if ovr.bold is not None:
                style.font.bold = ovr.bold
            if ovr.italic is not None:
                style.font.italic = ovr.italic
            pf = style.paragraph_format
            if ovr.space_before_pt is not None:
                pf.space_before = Pt(ovr.space_before_pt)
            if ovr.space_after_pt is not None:
                pf.space_after = Pt(ovr.space_after_pt)
            if ovr.line_spacing is not None:
                pf.line_spacing = ovr.line_spacing
            if ovr.line_spacing_twips is not None:
                pPr_elem = style.element.find(qn("w:pPr"))
                if pPr_elem is None:
                    pPr_elem = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                    style.element.append(pPr_elem)
                spacing = pPr_elem.find(qn("w:spacing"))
                if spacing is None:
                    spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
                    pPr_elem.append(spacing)
                spacing.set(qn("w:line"), str(ovr.line_spacing_twips))
                spacing.set(qn("w:lineRule"), "exact")
            if ovr.alignment:
                align = _alignment_enum(ovr.alignment)
                if align is not None:
                    pf.alignment = align
            if ovr.outline_level is not None:
                pPr_elem = style.element.find(qn("w:pPr"))
                if pPr_elem is None:
                    pPr_elem = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                    style.element.append(pPr_elem)
                outlineLvl = pPr_elem.find(qn("w:outlineLvl"))
                if outlineLvl is None:
                    outlineLvl = parse_xml(
                        f'<w:outlineLvl {nsdecls("w")} w:val="{ovr.outline_level}"/>'
                    )
                    pPr_elem.append(outlineLvl)
                else:
                    outlineLvl.set(qn("w:val"), str(ovr.outline_level))
        except Exception:
            pass


def _apply_word_instructions(doc: DocxDocument, instr: WordInstructions) -> None:
    if instr.default_font or instr.default_font_size_pt or instr.default_style:
        try:
            normal = doc.styles["Normal"]
            if instr.default_font:
                normal.font.name = instr.default_font
            if instr.default_font_size_pt:
                normal.font.size = Pt(instr.default_font_size_pt)
            elif instr.default_style:
                normal.font.name = instr.default_style
        except Exception:
            pass
    if instr.style_overrides:
        _apply_style_overrides(doc, instr.style_overrides)
    if instr.sections:
        for i, sec_spec in enumerate(instr.sections):
            if i == 0:
                _apply_section_props(doc.sections[0], sec_spec)
    if instr.content:
        for block in instr.content:
            if block.type == "paragraph" and block.paragraph:
                _add_paragraph(doc, block.paragraph)
            elif block.type == "table" and block.table:
                _add_table(doc, block.table)
            elif block.type == "image" and block.image:
                if block.image.b64:
                    data = base64.b64decode(block.image.b64)
                    kwargs = {"width": Inches(block.image.width_inches)}
                    if block.image.height_inches:
                        kwargs["height"] = Inches(block.image.height_inches)
                    doc.add_picture(io.BytesIO(data), **kwargs)
            elif block.type == "page_break":
                _add_page_break(doc)
            elif block.type == "section_break":
                _add_section_break(doc, block.section)
            elif block.type == "toc":
                _add_toc(doc, block.toc or TocSpec())
    else:
        for p in instr.paragraphs:
            _add_paragraph(doc, p)
        for t in instr.tables:
            _add_table(doc, t)
        for im in instr.images:
            if im.b64:
                data = base64.b64decode(im.b64)
                kwargs = {"width": Inches(im.width_inches)}
                if im.height_inches:
                    kwargs["height"] = Inches(im.height_inches)
                doc.add_picture(io.BytesIO(data), **kwargs)
    if instr.header_text:
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = instr.header_text
    if instr.footer_text:
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = instr.footer_text
    for fr in instr.find_replace:
        for para in doc.paragraphs:
            if fr.find in para.text:
                for run in para.runs:
                    if fr.find in run.text:
                        run.text = run.text.replace(fr.find, fr.replace)


# ============================================================
# Markdown Parsing & Table Helpers
# ============================================================


def _flush_table(
    doc,
    rows_raw: List[str],
    table_preset: str = "grid",
    default_font: Optional[str] = None,
    default_font_size_pt: Optional[float] = None,
) -> None:
    """Parse pipe-delimited table rows and add a formatted table."""
    parsed_rows = []
    for row_str in rows_raw:
        cells = [c.strip() for c in row_str.strip("|").split("|")]
        # Skip separator rows (e.g., |---|---|)
        if all(set(c) <= set("- :") for c in cells):
            continue
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    preset = table_preset or "grid"
    if preset == "none":
        _add_table(
            doc,
            TableSpec(
                rows=parsed_rows,
                header_bold=True,
                borders=False,
            ),
        )
    elif preset == "minimal":
        _add_table(
            doc,
            TableSpec(
                rows=parsed_rows,
                header_bold=True,
                borders=True,
                border_color="CCCCCC",
                border_size_pt=0.25,
            ),
        )
    else:  # "grid" (default)
        _add_table(
            doc,
            TableSpec(
                rows=parsed_rows,
                header_shading="D9D9D9",
                header_bold=True,
                borders=True,
                border_color="999999",
                border_size_pt=0.5,
            ),
        )


def _add_inline_formatted_runs(
    para,
    text: str,
    default_font: Optional[str] = None,
    default_font_size_pt: Optional[float] = None,
) -> None:
    """
    Parse inline markdown formatting (**bold**, *italic*, ***bold italic***)
    and add appropriately formatted runs to the paragraph.
    """
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)")
    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            run = para.add_run(text[last_end : match.start()])
            if default_font:
                run.font.name = default_font
            if default_font_size_pt:
                run.font.size = Pt(default_font_size_pt)

        if match.group(2):  # ***bold italic***
            run = para.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = para.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = para.add_run(match.group(4))
            run.italic = True

        if default_font:
            run.font.name = default_font
        if default_font_size_pt:
            run.font.size = Pt(default_font_size_pt)

        last_end = match.end()

    # Add remaining plain text
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        if default_font:
            run.font.name = default_font
        if default_font_size_pt:
            run.font.size = Pt(default_font_size_pt)


_NUMBERED_LIST_RE = re.compile(r"^(\d+)[\.\)]\s+(.+)$")


def _parse_markdown_content(
    doc,
    content: str,
    default_font: Optional[str] = None,
    default_font_size_pt: Optional[float] = None,
    table_preset: str = "grid",
) -> None:
    """
    Parse lightweight markdown content into Word elements.
    Supports: # H1, ## H2, ### H3, #### H4, - bullets, 1. numbered lists,
    **bold**, *italic*, | table rows |, and plain paragraphs.
    """
    if not content:
        return

    lines = content.split("\n")
    table_buffer: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            if table_buffer:
                _flush_table(
                    doc, table_buffer, table_preset, default_font, default_font_size_pt
                )
                table_buffer = []

        # Empty line -> skip
        if not stripped:
            i += 1
            continue

        # Heading 4
        if stripped.startswith("#### "):
            try:
                doc.add_paragraph(stripped[5:], style="Heading 4")
            except Exception:
                doc.add_paragraph(stripped[5:], style="Heading 3")
        # Heading 3
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 3")
        # Heading 2
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 2")
        # Heading 1
        elif stripped.startswith("# "):
            doc.add_paragraph(stripped[2:], style="Heading 1")
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            try:
                p = doc.add_paragraph(style="List Bullet")
            except Exception:
                p = doc.add_paragraph()
                bullet_text = "\u2022 " + bullet_text
            _add_inline_formatted_runs(
                p, bullet_text, default_font, default_font_size_pt
            )
        # Numbered list
        elif _NUMBERED_LIST_RE.match(stripped):
            m = _NUMBERED_LIST_RE.match(stripped)
            item_text = m.group(2)
            try:
                p = doc.add_paragraph(style="List Number")
            except Exception:
                p = doc.add_paragraph()
                item_text = m.group(1) + ". " + item_text
            _add_inline_formatted_runs(p, item_text, default_font, default_font_size_pt)
        # Normal paragraph with inline formatting
        else:
            p = doc.add_paragraph()
            _add_inline_formatted_runs(p, stripped, default_font, default_font_size_pt)

        i += 1

    # Flush remaining table
    if table_buffer:
        _flush_table(
            doc, table_buffer, table_preset, default_font, default_font_size_pt
        )


# ============================================================
# Document Builders
# ============================================================


def _build_general_document(meta: DocumentContent) -> bytes:
    """
    Build a general Word document from lightweight DocumentContent metadata.
    The tool handles all formatting internally.
    """
    doc = DocxDocument()

    # ---- Resolve style settings with defaults ----
    body_font = meta.font or "Calibri"
    body_size = meta.font_size_pt or 11.0
    h_font = meta.heading_font or body_font
    h_color = (meta.heading_color or "2F5496").lstrip("#")
    spacing = meta.line_spacing or 1.15
    margin = meta.margin_inches or 1.0
    table_preset = meta.table_style or "grid"

    # ---- Style overrides ----
    _apply_style_overrides(
        doc,
        [
            WordStyleOverride(
                style_id="Normal",
                font_name=body_font,
                font_size_pt=body_size,
                line_spacing=spacing,
                space_before_pt=2,
                space_after_pt=4,
            ),
            WordStyleOverride(
                style_id="Heading1",
                font_name=h_font,
                font_size_pt=16,
                bold=True,
                font_color=h_color,
                space_before_pt=14,
                space_after_pt=12,
                outline_level=0,
            ),
            WordStyleOverride(
                style_id="Heading2",
                font_name=h_font,
                font_size_pt=14,
                bold=True,
                font_color=h_color,
                space_before_pt=12,
                space_after_pt=6,
                outline_level=1,
            ),
            WordStyleOverride(
                style_id="Heading3",
                font_name=h_font,
                font_size_pt=12,
                bold=True,
                space_before_pt=10,
                space_after_pt=4,
                outline_level=2,
            ),
            WordStyleOverride(
                style_id="List Bullet",
                font_name=body_font,
                font_size_pt=body_size,
                line_spacing=spacing,
                space_before_pt=0,
                space_after_pt=2,
            ),
            WordStyleOverride(
                style_id="List Number",
                font_name=body_font,
                font_size_pt=body_size,
                line_spacing=spacing,
                space_before_pt=0,
                space_after_pt=2,
            ),
        ],
    )

    # ---- Page layout ----
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(margin)
    sec.bottom_margin = Inches(margin)
    sec.left_margin = Inches(margin)
    sec.right_margin = Inches(margin)

    if meta.orientation and meta.orientation.lower() == "landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width

    # ---- Header ----
    if meta.header_text:
        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(meta.header_text)
        hr.font.size = Pt(9)
        hr.font.name = body_font
        hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ---- Footer ----
    if meta.footer_text or meta.include_page_numbers:
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if meta.footer_text:
            fr = fp.add_run(meta.footer_text + "   ")
            fr.font.size = Pt(9)
            fr.font.name = body_font
            fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        if meta.include_page_numbers:
            run1 = fp.add_run()
            run1._element.append(
                parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            )
            run2 = fp.add_run()
            run2._element.append(
                parse_xml(
                    f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
                )
            )
            run3 = fp.add_run()
            run3._element.append(
                parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            )

    # ---- Title ----
    if meta.title:
        p = doc.add_paragraph()
        r = p.add_run(meta.title)
        r.bold = True
        r.font.size = Pt(22)
        r.font.name = h_font
        r.font.color.rgb = _hex_to_rgb(h_color)
        p.paragraph_format.space_after = Pt(6)

    # ---- Subtitle ----
    if meta.subtitle:
        p = doc.add_paragraph()
        r = p.add_run(meta.subtitle)
        r.font.size = Pt(10)
        r.font.name = body_font
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        p.paragraph_format.space_after = Pt(12)

    # ---- Table of Contents ----
    if meta.include_toc:
        _add_toc(doc, TocSpec(title="Table of Contents", max_level=3))
        _add_page_break(doc)

    # ---- Body sections ----
    for section_data in meta.sections:
        if section_data.heading:
            doc.add_paragraph(section_data.heading, style="Heading 1")
        _parse_markdown_content(
            doc,
            section_data.content,
            default_font=body_font,
            default_font_size_pt=body_size,
            table_preset=table_preset,
        )

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _create_docx(instr: WordInstructions) -> bytes:
    doc = DocxDocument()
    _apply_word_instructions(doc, instr)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _modify_docx(existing_bytes: bytes, instr: WordInstructions) -> bytes:
    bio = io.BytesIO(existing_bytes)
    doc = DocxDocument(bio)
    _apply_word_instructions(doc, instr)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ============================================================
# Public API
# ============================================================


def generate_docx(
    content: DocumentContent,
    output_path: Optional[str] = None,
) -> str:
    """
    Generate a .docx file from a DocumentContent specification.

    Args:
        content: DocumentContent with markdown sections and optional style hints.
        output_path: Where to save the .docx. If None, auto-generates a name in cwd.

    Returns:
        Absolute path to the generated .docx file.

    Example:
        >>> from word_doc_gen import generate_docx, DocumentContent, DocumentSection
        >>> doc = DocumentContent(
        ...     title="My Resume",
        ...     font="Calibri",
        ...     font_size_pt=11,
        ...     sections=[
        ...         DocumentSection(heading="Experience", content="- Job 1\\n- Job 2")
        ...     ]
        ... )
        >>> path = generate_docx(doc, "/tmp/resume.docx")
    """
    data = _build_general_document(content)
    output_path = output_path or _choose_output_name(content.title or None)
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "wb") as f:
        f.write(data)
    return output_path


def generate_docx_from_instructions(
    instructions: WordInstructions,
    output_path: Optional[str] = None,
    existing_docx_bytes: Optional[bytes] = None,
) -> str:
    """
    Generate or modify a .docx file from full WordInstructions.

    Args:
        instructions: Full WordInstructions specification.
        output_path: Where to save the .docx. Auto-generated if None.
        existing_docx_bytes: If provided, modifies this existing document
                             instead of creating from scratch.

    Returns:
        Absolute path to the generated .docx file.
    """
    if existing_docx_bytes:
        data = _modify_docx(existing_docx_bytes, instructions)
    else:
        data = _create_docx(instructions)
    output_path = output_path or _choose_output_name(None)
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "wb") as f:
        f.write(data)
    return output_path


# ============================================================
# Resume Builder
# ============================================================

# Markdown section headings that trigger resume-specific parsing
_RESUME_SECTION_TYPES = {
    "professional summary": "summary",
    "technical skills": "skills",
    "certifications": "certifications",
    "professional experience": "experience",
    "research and projects": "projects",
    "research & projects": "projects",
    "education": "education",
    "professional development": "development",
}

# Lines that look like skill labels (e.g., "Offensive Security: ...")
_SKILL_LABEL_RE = re.compile(
    r"^(\*\*)?(.+?):(\*\*)?\s+(.+)$"
)

# Role title lines (### Header or bold text followed by company)
_ROLE_HEADER_RE = re.compile(r"^###\s+(.+)$")

# Company/date lines: **Company** | Date
_COMPANY_DATE_RE = re.compile(r"^\*\*(.+?)\*\*\s*\|(.+)$")

# Project bullet with bold name prefix: **name:** or **name (**
_PROJECT_NAME_RE = re.compile(r"^\*\*(.+?)\*\*\s*[\(:]\s*(.+)$")


def _setup_resume_defaults(doc: DocxDocument, meta: ResumeContent) -> None:
    """Configure Normal style and page layout to match the resume template."""
    # Override Normal style
    try:
        normal = doc.styles["Normal"]
        normal.font.name = meta.font_family
        normal.font.size = Pt(meta.font_size_pt)
        normal.font.color.rgb = _hex_to_rgb("1E1E1E")
        normal.paragraph_format.space_after = Pt(0)
        # Set exact line spacing (240 twips = 1.0x for 9pt)
        pPr = normal.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            normal.element.append(pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(spacing)
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:after"), "0")
    except Exception:
        pass

    # Create ResumeSection custom style
    try:
        rs_style = doc.styles.add_style("Resume Section", 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
        rs_style.font.name = meta.font_family
        rs_style.font.size = Pt(10.5)
        rs_style.font.bold = True
        rs_style.font.color.rgb = _hex_to_rgb(meta.header_color)
        rs_pPr = rs_style.element.find(qn("w:pPr"))
        if rs_pPr is None:
            rs_pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            rs_style.element.append(rs_pPr)
        # Spacing: before 160, after 60
        rs_spacing = rs_pPr.find(qn("w:spacing"))
        if rs_spacing is None:
            rs_spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            rs_pPr.append(rs_spacing)
        rs_spacing.set(qn("w:before"), "160")
        rs_spacing.set(qn("w:after"), "60")
        # Bottom border
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="2" w:color="{meta.border_color}"/>'
            f'</w:pBdr>'
        )
        rs_pPr.append(pBdr)
    except Exception:
        pass

    # Override List Bullet style
    try:
        lb = doc.styles["List Bullet"]
        lb.font.name = meta.font_family
        lb.font.size = Pt(meta.font_size_pt)
        pPr = lb.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            lb.element.append(pPr)
    except Exception:
        pass

    # Page layout: Letter, matching margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)


def _add_resume_section_header(doc: DocxDocument, text: str) -> None:
    """Add a resume section header with the custom Resume Section style and border."""
    para = doc.add_paragraph()
    try:
        para.style = doc.styles["Resume Section"]
    except Exception:
        pass
    para.paragraph_format.keep_with_next = True
    # Ensure bottom border is on the paragraph element
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="2" w:color="D9E2F3"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    # Set line spacing to 240 auto
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    # Add text as run with the style's formatting
    run = para.add_run(text.upper())


def _add_resume_name_header(doc: DocxDocument, name: str) -> None:
    """Add the name: centered, bold, 20pt, #245291."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(name)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = _hex_to_rgb("245291")


def _add_resume_subtitle(doc: DocxDocument, text: str) -> None:
    """Add the subtitle line: centered, bold. Strips markdown bold markers."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(1.4)  # ~20 twips
    # Strip ** markers from plain text
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    run = para.add_run(clean)
    run.bold = True


def _add_resume_contact(doc: DocxDocument, text: str) -> None:
    """Add the contact info line: centered."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(text)


def _add_summary_body(doc: DocxDocument, text: str) -> None:
    """Add the professional summary paragraph."""
    para = doc.add_paragraph()
    para.add_run(text)
    # Set line spacing to 247 auto
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn("w:line"), "247")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:after"), "60")


def _add_skill_line(doc: DocxDocument, line: str) -> None:
    """Add a skill line: bold prefix label, rest normal."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2.1)  # ~30 twips
    m = _SKILL_LABEL_RE.match(line.strip())
    if m:
        label = m.group(2).strip()
        rest = m.group(4).strip()
        run_label = para.add_run(label + ": ")
        run_label.bold = True
        run_rest = para.add_run(rest)
    else:
        run = para.add_run(line.strip())
        run.bold = True


def _add_role_title(doc: DocxDocument, text: str) -> None:
    """Add a role title: bold, 10pt, space before 80, keep with next."""
    para = doc.add_paragraph()
    para.paragraph_format.keep_with_next = True
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn("w:before"), "80")
    run = para.add_run(text.strip())
    run.bold = True
    run.font.size = Pt(10)


def _add_company_date(doc: DocxDocument, line: str) -> None:
    """Add company and dates: company bold, dates gray #505050."""
    para = doc.add_paragraph()
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.space_after = Pt(1.4)  # ~20 twips
    m = _COMPANY_DATE_RE.match(line.strip())
    if m:
        company = m.group(1).strip()
        dates = m.group(2).strip()
        run_company = para.add_run(company)
        run_company.bold = True
        run_dates = para.add_run(" | " + dates)
        run_dates.font.color.rgb = _hex_to_rgb("505050")
    else:
        run = para.add_run(line.strip())
        run.bold = True


def _add_experience_bullet(doc: DocxDocument, text: str, use_compact: bool = False) -> None:
    """Add a bullet list item under Professional Experience."""
    try:
        para = doc.add_paragraph(style="List Bullet")
    except Exception:
        para = doc.add_paragraph()
    if use_compact:
        para.paragraph_format.space_after = Pt(2.1)  # ~30 twips
        pPr = para._element.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
            pPr.append(ind)
        ind.set(qn("w:left"), "331")
        ind.set(qn("w:hanging"), "187")
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(spacing)
        spacing.set(qn("w:line"), "247")
        spacing.set(qn("w:lineRule"), "auto")
    para.add_run(text.strip())


def _add_project_bullet(doc: DocxDocument, line: str) -> None:
    """Add a project bullet with bold name prefix."""
    try:
        para = doc.add_paragraph(style="List Bullet")
    except Exception:
        para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2.1)  # ~30 twips
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
        pPr.append(ind)
    ind.set(qn("w:left"), "331")
    ind.set(qn("w:hanging"), "187")
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn("w:line"), "245")
    spacing.set(qn("w:lineRule"), "auto")

    stripped = line.strip()
    m = _PROJECT_NAME_RE.match(stripped)
    if m:
        name = m.group(1).strip()
        rest = m.group(2).strip()
        run_name = para.add_run(name + ": ")
        run_name.bold = True
        run_name.font.size = Pt(8.5)
        run_rest = para.add_run(rest)
        run_rest.font.size = Pt(8.5)
    else:
        run = para.add_run(stripped)
        run.font.size = Pt(8.5)


def _add_education_line(doc: DocxDocument, text: str) -> None:
    """Add an education line: 8.5pt, after 20, line 247."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(1.4)  # ~20 twips
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn("w:line"), "247")
    spacing.set(qn("w:lineRule"), "auto")
    run = para.add_run(text.strip())
    run.font.size = Pt(8.5)


def _add_cert_paragraph(doc: DocxDocument, text: str) -> None:
    """Add a certifications paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2.8)  # ~40 twips
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn("w:line"), "247")
    spacing.set(qn("w:lineRule"), "auto")
    para.add_run(text.strip())


def _parse_resume_section(
    doc: DocxDocument,
    heading: str,
    content: str,
    meta: ResumeContent,
) -> None:
    """Parse and render a single resume section by type."""
    section_type = _RESUME_SECTION_TYPES.get(heading.lower().strip(), "generic")
    lines = [l for l in content.split("\n") if l.strip()]

    if section_type == "summary":
        text = " ".join(lines)
        _add_summary_body(doc, text)
        return

    if section_type == "skills":
        for line in lines:
            _add_skill_line(doc, line)
        return

    if section_type == "certifications":
        for line in lines:
            _add_cert_paragraph(doc, line)
        return

    if section_type == "experience":
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith("### "):
                title = line[4:].strip()
                _add_role_title(doc, title)
                i += 1
                if i < len(lines):
                    next_line = lines[i].strip()
                    if _COMPANY_DATE_RE.match(next_line):
                        _add_company_date(doc, next_line)
                        i += 1
                continue
            elif _COMPANY_DATE_RE.match(line):
                _add_company_date(doc, line)
                i += 1
                continue
            elif line.startswith("*") or line.startswith("-"):
                bullet_text = re.sub(r"^[*\-]\s+", "", line).strip()
                _add_experience_bullet(doc, bullet_text)
                i += 1
                continue
            elif line:
                if i + 1 < len(lines) and _COMPANY_DATE_RE.match(lines[i + 1].strip()):
                    _add_role_title(doc, line)
                    i += 1
                    _add_company_date(doc, lines[i].strip())
                    i += 1
                    continue
                para = doc.add_paragraph(line)
                i += 1
                continue
            i += 1
        return

    if section_type == "projects":
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("*") or stripped.startswith("-"):
                bullet_text = re.sub(r"^[*\-]\s+", "", stripped)
                _add_project_bullet(doc, bullet_text)
            else:
                _add_project_bullet(doc, stripped)
        return

    if section_type == "education":
        for line in lines:
            _add_education_line(doc, line)
        return

    if section_type == "development":
        for line in lines:
            _add_skill_line(doc, line)
        return

    # Generic fallback
    for line in lines:
        doc.add_paragraph(line)


def _parse_markdown_resume(
    doc: DocxDocument,
    markdown: str,
    meta: ResumeContent,
) -> None:
    """Parse the full markdown resume and build the document."""
    lines = markdown.split("\n")

    # Phase 1: Extract name (H1), subtitle lines, and contact line
    i = 0
    name = ""
    subtitle_lines: List[str] = []
    contact = ""

    while i < len(lines) and not lines[i].strip():
        i += 1

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# ") and not stripped.startswith("## "):
            name = stripped[2:].strip()
            name = re.sub(r"\s*[—–-]\s*Resume\s*$", "", name)
            i += 1
            break
        elif not stripped.startswith("#") and not stripped.startswith(">"):
            name = stripped
            i += 1
            break
        else:
            i += 1
            break

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("## ") and not stripped.startswith("### "):
            break
        if stripped.startswith(">"):
            i += 1
            continue
        # Detect contact line (contains email, linkedin, or github)
        is_contact = (
            "@" in stripped or "linkedin" in stripped.lower() or "github" in stripped.lower()
        )
        # Also detect contact by common patterns: phone numbers, location markers
        if not is_contact:
            has_phone = bool(re.search(r'\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}', stripped))
            has_separators = stripped.count("\u2022") >= 2 or stripped.count("|") >= 3
            is_contact = has_phone or has_separators

        if is_contact:
            if contact:
                # Already have contact — join with " | "
                contact += " | " + stripped
            else:
                contact = stripped
        elif not contact:
            subtitle_lines.append(stripped)
        else:
            subtitle_lines.append(stripped)
        i += 1

    if subtitle_lines and not contact:
        last = subtitle_lines[-1]
        has_phone = bool(re.search(r'\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}', last))
        has_separators = last.count("\u2022") >= 2 or last.count("|") >= 3
        if "@" in last or "linkedin" in last.lower() or has_phone or has_separators:
            contact = last
            subtitle_lines = subtitle_lines[:-1]

    # Render header block
    if name:
        _add_resume_name_header(doc, name)
    if subtitle_lines:
        _add_resume_subtitle(doc, " | ".join(subtitle_lines))
    if contact:
        _add_resume_contact(doc, contact)

    # Phase 2: Parse H2 sections
    current_heading = ""
    current_content: List[str] = []

    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith("## ") and not stripped.startswith("### "):
            if current_heading:
                _add_resume_section_header(doc, current_heading)
                _parse_resume_section(
                    doc, current_heading, "\n".join(current_content), meta,
                )
            current_heading = stripped[3:].strip()
            current_content = []
        elif current_heading:
            current_content.append(lines[i])
        i += 1

    if current_heading:
        _add_resume_section_header(doc, current_heading)
        _parse_resume_section(
            doc, current_heading, "\n".join(current_content), meta,
        )


def _build_resume_document(meta: ResumeContent) -> bytes:
    """Build a resume .docx from markdown content with exact template formatting."""
    doc = DocxDocument()
    _setup_resume_defaults(doc, meta)
    _parse_markdown_resume(doc, meta.markdown, meta)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_resume_docx(
    markdown_path: str,
    output_path: Optional[str] = None,
) -> str:
    """
    Generate a resume .docx from a markdown resume file.

    Uses the canonical resume template formatting:
    - Letter paper, 0.5"/0.45"/0.6"/0.6" margins
    - Arial 9pt body, 10.5pt section headers with #245291 color
    - Resume Section style with #D9E2F3 bottom border
    - Centered name header (20pt bold, #245291)
    - Role titles (10pt bold), company/dates (company bold, dates gray)
    - Bullet lists with Symbol bullet character

    Args:
        markdown_path: Path to the markdown resume file.
        output_path: Where to save the .docx. Defaults to
            .pi/skills/rez/output/Kristoffer_Sketch_Resume.docx.

    Returns:
        Absolute path to the generated .docx file.
    """
    if not os.path.isfile(markdown_path):
        raise FileNotFoundError(f"Markdown resume not found: {markdown_path}")

    with open(markdown_path, "r") as f:
        markdown_content = f.read()

    if not markdown_content.strip():
        raise ValueError("Markdown file is empty")

    if output_path is None:
        output_path = os.path.join(
            os.path.dirname(__file__), "..", "..", ".pi", "skills", "resume",
            "output", "Kristoffer_Sketch_Resume.docx"
        )

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    content = ResumeContent(markdown=markdown_content)
    data = _build_resume_document(content)

    with open(output_path, "wb") as f:
        f.write(data)

    return output_path
