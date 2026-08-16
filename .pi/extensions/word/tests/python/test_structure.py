from __future__ import annotations

from pathlib import Path
from types import ModuleType

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image


def _generate(
    tmp_path: Path, wordgen: ModuleType, name: str, markdown: str, **options: object
) -> DocxDocument:
    output = tmp_path / f"{name}.docx"
    wordgen.generate(
        {
            "markdown": markdown,
            "title_mode": "none",
            "output_path": str(output),
            "project_root": str(tmp_path),
            **options,
        }
    )
    return Document(str(output))


def test_ordered_list_continuation_has_no_bullet_or_number(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "ordered-continuation",
        "1. First paragraph\n\n   Continuation paragraph\n2. Second item",
    )
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "1.\tFirst paragraph",
        "Continuation paragraph",
        "2.\tSecond item",
    ]
    continuation = document.paragraphs[1]
    assert continuation.style is not None
    assert continuation.style.name == wordgen.PENNY_LIST_CONTINUE
    assert continuation._p.xpath(".//w:numPr") == []


def test_separate_ordered_lists_restart_at_one(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "list-restart",
        "1. Alpha\n2. Beta\n\nParagraph.\n\n1. Gamma",
    )
    numbered = [paragraph.text for paragraph in document.paragraphs if "\t" in paragraph.text]
    assert numbered == ["1.\tAlpha", "2.\tBeta", "1.\tGamma"]


def test_content_layout_allocates_more_width_to_long_column(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "content-table",
        "| Status | Description |\n"
        "| --- | --- |\n"
        "| OK | This is a substantially longer explanatory value for the second column. |",
    )
    grid_columns = document.tables[0]._tbl.tblGrid.gridCol_lst
    assert grid_columns[0].w is not None
    assert grid_columns[1].w is not None
    assert grid_columns[1].w > grid_columns[0].w


def test_equal_layout_preserves_equal_columns(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "equal-table",
        "| Status | Description |\n| --- | --- |\n| OK | A much longer value. |",
        table_layout="equal",
    )
    grid_columns = document.tables[0]._tbl.tblGrid.gridCol_lst
    assert grid_columns[0].w == grid_columns[1].w


def test_markdown_alignment_maps_to_word_alignment(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "aligned-table",
        "| Left | Center | Right |\n| :--- | :---: | ---: |\n| A | B | C |",
    )
    row = document.tables[0].rows[1]
    assert row.cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert row.cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert row.cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_table_cells_have_compact_spacing_and_row_policies(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "table-policy",
        "| Header | Value |\n| --- | --- |\n| A | B |",
    )
    table = document.tables[0]
    for row in table.rows:
        assert len(row._tr.xpath("./w:trPr/w:cantSplit")) == 1
        for cell in row.cells:
            formatting = cell.paragraphs[0].paragraph_format
            assert formatting.space_before.pt == 0
            assert formatting.space_after.pt == 0
    assert len(table.rows[0]._tr.xpath("./w:trPr/w:tblHeader")) == 1


def test_missing_image_warns_without_corrupting_document(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    output = tmp_path / "missing-image.docx"
    result = wordgen.generate(
        {
            "markdown": "![missing image](not-present.png)",
            "title_mode": "none",
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )
    assert result["warnings"] == ["image not found: not-present.png"]
    assert "image unavailable" in Document(str(output)).paragraphs[0].text


def test_block_figure_keeps_image_with_caption(tmp_path: Path, wordgen: ModuleType) -> None:
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (100, 60), color=(0, 100, 180)).save(image_path)
    document = _generate(
        tmp_path,
        wordgen,
        "block-figure",
        "![Figure caption](figure.png)",
    )
    image_paragraph, caption = document.paragraphs
    assert image_paragraph.paragraph_format.keep_with_next is True
    assert caption.paragraph_format.keep_with_next is False
    assert caption.text == "Figure caption"


def test_inline_image_stays_in_authored_paragraph_and_has_alt_text(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    image_path = tmp_path / "dot.png"
    Image.new("RGB", (20, 20), color=(255, 0, 0)).save(image_path)
    document = _generate(
        tmp_path,
        wordgen,
        "inline-image",
        "Before ![status dot](dot.png) after.",
    )
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "Before  after."
    assert len(document.paragraphs[0]._p.xpath(".//w:drawing")) == 1
    assert document.inline_shapes[0]._inline.docPr.get("descr") == "status dot"


def test_code_block_is_one_continuous_paragraph(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "code-block",
        "```python\nfirst = 1\nsecond = 2\n```",
    )
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].style is not None
    assert document.paragraphs[0].style.name == wordgen.PENNY_CODE_BLOCK
    assert document.paragraphs[0].text == "first = 1\nsecond = 2"
    assert len(document.paragraphs[0]._p.xpath(".//w:br")) == 1


def test_cover_isolated_from_body_footer_and_restarts_page_number(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    output = tmp_path / "cover.docx"
    wordgen.generate(
        {
            "markdown": "# Report\n\n## Body\n\nContent.",
            "title_mode": "cover",
            "footer_text": "Confidential",
            "include_page_numbers": True,
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )
    document = Document(str(output))
    assert len(document.sections) == 2
    cover, body = document.sections
    assert cover.footer.is_linked_to_previous is True
    assert body.footer.is_linked_to_previous is False
    page_number = body._sectPr.find(qn("w:pgNumType"))
    assert page_number is not None
    assert page_number.get(qn("w:start")) == "1"


def test_toc_requests_field_refresh_and_reports_its_limit(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    output = tmp_path / "toc.docx"
    result = wordgen.generate(
        {
            "markdown": "# Report\n\n## Body\n\nContent.",
            "include_toc": True,
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )
    document = Document(str(output))
    update_fields = document.settings._element.find(qn("w:updateFields"))
    assert update_fields is not None
    assert update_fields.get(qn("w:val")) == "true"
    assert result["toc_field_update_requested"] is True
    assert any("TOC field" in warning for warning in result["warnings"])
