from __future__ import annotations

from pathlib import Path
from types import ModuleType

import pytest
from docx import Document
from docx.document import Document as DocxDocument


def _generate(
    tmp_path: Path, wordgen: ModuleType, markdown: str, **options: object
) -> DocxDocument:
    output = tmp_path / f"break-{len(list(tmp_path.iterdir()))}.docx"
    wordgen.generate(
        {
            "markdown": markdown,
            "title_mode": "none",
            "output_path": str(output),
            "project_root": str(tmp_path),
            **options,
        }
    )
    return Document(str(output))


def _break_count(paragraph: object) -> int:
    return len(paragraph._p.xpath(".//w:br"))  # type: ignore[attr-defined]


def test_preserve_mode_turns_softbreak_into_word_break(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(tmp_path, wordgen, "Alpha\nBeta", line_break_mode="preserve")
    assert document.paragraphs[0].text == "Alpha\nBeta"
    assert _break_count(document.paragraphs[0]) == 1


def test_commonmark_mode_turns_softbreak_into_space(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(tmp_path, wordgen, "Alpha\nBeta", line_break_mode="commonmark")
    assert document.paragraphs[0].text == "Alpha Beta"
    assert _break_count(document.paragraphs[0]) == 0


def test_blank_line_creates_new_paragraph_in_both_modes(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    for mode in ("preserve", "commonmark"):
        document = _generate(tmp_path, wordgen, "Alpha\n\nBeta", line_break_mode=mode)
        assert [paragraph.text for paragraph in document.paragraphs] == ["Alpha", "Beta"]


def test_hardbreak_creates_word_break_in_both_modes(tmp_path: Path, wordgen: ModuleType) -> None:
    for mode in ("preserve", "commonmark"):
        document = _generate(tmp_path, wordgen, "Alpha  \nBeta", line_break_mode=mode)
        assert document.paragraphs[0].text == "Alpha\nBeta"
        assert _break_count(document.paragraphs[0]) == 1


@pytest.mark.parametrize("tag", ["<br>", "<br/>", "<br />", "<BR>"])
def test_html_br_variants_create_word_break(tmp_path: Path, wordgen: ModuleType, tag: str) -> None:
    document = _generate(tmp_path, wordgen, f"Alpha{tag}Beta", line_break_mode="commonmark")
    assert document.paragraphs[0].text == "Alpha\nBeta"
    assert _break_count(document.paragraphs[0]) == 1


def test_break_inside_hyperlink_remains_inside_hyperlink(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "[Alpha\nBeta](https://example.test)",
        line_break_mode="preserve",
    )
    hyperlinks = document.paragraphs[0]._p.xpath(".//w:hyperlink")
    assert len(hyperlinks) == 1
    assert len(hyperlinks[0].xpath(".//w:br")) == 1


def test_break_policy_applies_in_heading(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(tmp_path, wordgen, "## Alpha<br>Beta", line_break_mode="preserve")
    assert document.paragraphs[0].style is not None
    assert document.paragraphs[0].style.name == "Heading 2"
    assert _break_count(document.paragraphs[0]) == 1


def test_softbreak_policy_applies_in_ordered_list_item(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(tmp_path, wordgen, "1. Alpha\n   Beta", line_break_mode="preserve")
    assert document.paragraphs[0].text == "1.\tAlpha\nBeta"
    assert _break_count(document.paragraphs[0]) == 1


def test_softbreak_policy_applies_in_table_cell(tmp_path: Path, wordgen: ModuleType) -> None:
    document = _generate(
        tmp_path,
        wordgen,
        "| Value |\n| --- |\n| Alpha<br>Beta |",
        line_break_mode="commonmark",
    )
    cell_paragraph = document.tables[0].cell(1, 0).paragraphs[0]
    assert cell_paragraph.text == "Alpha\nBeta"
    assert _break_count(cell_paragraph) == 1
