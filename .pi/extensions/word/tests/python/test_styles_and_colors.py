from __future__ import annotations

from pathlib import Path
from types import ModuleType
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _options(tmp_path: Path, wordgen: ModuleType, **overrides: object) -> Any:
    return wordgen.parse_options(
        {
            "output_path": str(tmp_path / "unused.docx"),
            "project_root": str(tmp_path),
            **overrides,
        }
    )


def test_heading_scale_is_valid_for_every_theme_and_body_extreme(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    for theme in wordgen.THEMES:
        for body_size in (8, 11, 14):
            options = _options(tmp_path, wordgen, theme=theme, font_size_pt=body_size)
            scale = options.scale
            assert scale.title > scale.h1 > scale.h2 > scale.h3 >= scale.body
            assert scale.h4 >= scale.h5 >= scale.h6


def test_h4_h5_h6_have_distinct_style_tuples(tmp_path: Path, wordgen: ModuleType) -> None:
    options = _options(tmp_path, wordgen, font_size_pt=14)
    document = Document()
    wordgen._setup_styles(document, options)

    tuples = []
    for name in ("Heading 4", "Heading 5", "Heading 6"):
        font = document.styles[name].font
        tuples.append((font.size.pt, font.bold, font.italic, str(font.color.rgb)))
    assert len(set(tuples)) == 3


def test_table_caption_code_header_footer_sizes_derive_from_body_scale(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    small = _options(tmp_path, wordgen, font_size_pt=8).scale
    large = _options(tmp_path, wordgen, font_size_pt=14).scale
    for attribute in ("table", "caption", "code", "header_footer"):
        assert getattr(large, attribute) > getattr(small, attribute)


def test_ordinary_body_runs_inherit_font_size_and_color_from_style(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    output = tmp_path / "inherit.docx"
    wordgen.generate(
        {
            "markdown": "Ordinary body text.",
            "title_mode": "none",
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )
    paragraph = Document(str(output)).paragraphs[0]
    assert paragraph.style is not None
    assert paragraph.style.name == wordgen.PENNY_BODY
    run_properties = paragraph.runs[0]._r.rPr
    assert run_properties is not None
    assert run_properties.find(qn("w:rFonts")) is None
    assert run_properties.find(qn("w:sz")) is None
    assert run_properties.find(qn("w:color")) is None


def test_bold_italic_strike_hyperlink_and_code_semantics_survive(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    output = tmp_path / "semantics.docx"
    wordgen.generate(
        {
            "markdown": "**bold** *italic* ~~strike~~ `code` [link](https://example.test)",
            "title_mode": "none",
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )
    paragraph = Document(str(output)).paragraphs[0]
    assert any(run.text == "bold" and run.bold for run in paragraph.runs)
    assert any(run.text == "italic" and run.italic for run in paragraph.runs)
    assert any(run.text == "strike" and run.font.strike for run in paragraph.runs)
    assert any(
        run.text == "code" and run.style.name == wordgen.PENNY_INLINE_CODE for run in paragraph.runs
    )
    hyperlink_styles = paragraph._p.xpath(".//w:hyperlink/w:r/w:rPr/w:rStyle/@w:val")
    assert hyperlink_styles == ["PennyHyperlink"]


def test_hyperlinks_inherit_heading_and_unfilled_table_typography(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    output = tmp_path / "contextual-links.docx"
    wordgen.generate(
        {
            "markdown": (
                "# [Linked Heading](https://example.test)\n\n"
                "| [Header](https://example.test) |\n| --- |\n| Value |"
            ),
            "title_mode": "none",
            "table_style": "minimal",
            "font_size_pt": 14,
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )
    document = Document(str(output))
    hyperlink_style = document.styles[wordgen.PENNY_HYPERLINK]
    assert hyperlink_style.font.name is None
    assert hyperlink_style.font.size is None
    assert document.paragraphs[0].style is not None
    assert document.paragraphs[0].style.font.size.pt == document.styles["Heading 1"].font.size.pt
    table_header = document.tables[0].cell(0, 0).paragraphs[0]
    assert table_header.style is not None
    assert (
        table_header.style.font.size.pt == document.styles[wordgen.PENNY_TABLE_HEADER].font.size.pt
    )


def test_inline_code_style_inherits_size_but_sets_monospace_font(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    options = _options(tmp_path, wordgen)
    document = Document()
    wordgen._setup_styles(document, options)
    style = document.styles[wordgen.PENNY_INLINE_CODE]
    assert style.font.name == options.theme.mono_font
    assert style.font.size is None


def test_named_styles_populate_all_script_font_slots(tmp_path: Path, wordgen: ModuleType) -> None:
    options = _options(tmp_path, wordgen)
    document = Document()
    wordgen._setup_styles(document, options)

    for style_name in (
        wordgen.PENNY_BODY,
        "Heading 1",
        wordgen.PENNY_TABLE_BODY,
        wordgen.PENNY_INLINE_CODE,
    ):
        rfonts = document.styles[style_name]._element.get_or_add_rPr().find(qn("w:rFonts"))
        assert rfonts is not None
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            assert rfonts.get(qn(f"w:{slot}"))


def test_custom_accent_regenerates_soft_palette(tmp_path: Path, wordgen: ModuleType) -> None:
    default = _options(tmp_path, wordgen, theme="executive").colors
    custom = _options(tmp_path, wordgen, theme="executive", accent_color="FDE047").colors
    assert custom.accent == "FDE047"
    assert custom.accent_soft != default.accent_soft
    assert custom.accent_soft == wordgen._mix("FDE047", "FFFFFF", 0.86)


def test_text_role_pairs_meet_contrast_threshold(tmp_path: Path, wordgen: ModuleType) -> None:
    for accent in ("FDE047", "111827", "777777", "FF00FF", "00FFFF"):
        colors = _options(tmp_path, wordgen, accent_color=accent).colors
        assert wordgen.contrast_ratio(colors.text, colors.background) >= 4.5
        assert wordgen.contrast_ratio(colors.text_muted, colors.background) >= 4.5
        assert wordgen.contrast_ratio(colors.heading, colors.background) >= 4.5
        assert wordgen.contrast_ratio(colors.link, colors.background) >= 4.5
        assert wordgen.contrast_ratio(colors.on_accent, colors.accent) >= 4.5
        assert wordgen.contrast_ratio(colors.link_on_accent, colors.accent) >= 4.5


def test_table_header_link_uses_on_accent_style(tmp_path: Path, wordgen: ModuleType) -> None:
    output = tmp_path / "header-link.docx"
    result = wordgen.generate(
        {
            "markdown": "| [Header](https://example.test) |\n| --- |\n| Value |",
            "title_mode": "none",
            "accent_color": "FDE047",
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )
    paragraph = Document(str(output)).tables[0].cell(0, 0).paragraphs[0]
    style_values = paragraph._p.xpath(".//w:hyperlink/w:r/w:rPr/w:rStyle/@w:val")
    assert style_values == ["PennyHyperlinkOnAccent"]
    palette = result["resolved_palette"]
    assert wordgen.contrast_ratio(palette["link_on_accent"], palette["accent"]) >= 4.5
