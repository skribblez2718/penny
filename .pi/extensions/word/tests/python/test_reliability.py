from __future__ import annotations

import json
import subprocess
import sys
import zipfile
from pathlib import Path
from types import ModuleType

import pytest
from docx import Document

GENERATOR = Path(__file__).parents[2] / "generate_docx.py"
REQUIRED_PARTS = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}


def test_generate_minimal_docx_has_required_parts_crc_and_reopens(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    output = tmp_path / "minimal.docx"
    result = wordgen.generate(
        {
            "markdown": "# Hello\n\nA **bold** paragraph with café and 東京.",
            "title_mode": "none",
            "output_path": str(output),
            "project_root": str(tmp_path),
        }
    )

    assert Path(result["path"]) == output
    assert result["validation"]["package_valid"] is True
    assert result["validation"]["reopen_valid"] is True
    with zipfile.ZipFile(output) as archive:
        assert REQUIRED_PARTS <= set(archive.namelist())
        assert archive.testzip() is None
    reopened = Document(str(output))
    assert "café" in "\n".join(paragraph.text for paragraph in reopened.paragraphs)


def test_validation_rejects_non_zip_package(tmp_path: Path, wordgen: ModuleType) -> None:
    invalid = tmp_path / "invalid.docx"
    invalid.write_text("not a zip", encoding="utf-8")
    with pytest.raises(ValueError, match="invalid generated DOCX"):
        wordgen._validate_docx(invalid)


def test_validation_rejects_missing_required_part(tmp_path: Path, wordgen: ModuleType) -> None:
    invalid = tmp_path / "missing.docx"
    with zipfile.ZipFile(invalid, "w") as archive:
        archive.writestr("word/document.xml", "<root/>")
    with pytest.raises(ValueError, match="missing required parts"):
        wordgen._validate_docx(invalid)


def test_validation_rejects_malformed_xml_part(tmp_path: Path, wordgen: ModuleType) -> None:
    invalid = tmp_path / "malformed.docx"
    with zipfile.ZipFile(invalid, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types>")
        archive.writestr("_rels/.rels", "<Relationships/>")
        archive.writestr("word/document.xml", "<document/>")
    with pytest.raises(ValueError, match="invalid generated DOCX"):
        wordgen._validate_docx(invalid)


def test_atomic_failure_preserves_existing_target_bytes(
    tmp_path: Path, wordgen: ModuleType, monkeypatch: pytest.MonkeyPatch
) -> None:
    target = tmp_path / "existing.docx"
    original = b"known-good-prior-file"
    target.write_bytes(original)
    document = Document()
    document.add_paragraph("replacement")

    def reject_validation(_path: Path) -> None:
        raise ValueError("forced validation failure")

    monkeypatch.setattr(wordgen, "_validate_docx", reject_validation)
    with pytest.raises(ValueError, match="forced validation failure"):
        wordgen._save_docx_atomically(document, target)

    assert target.read_bytes() == original
    assert list(tmp_path.glob(".*.tmp.docx")) == []


def test_atomic_failure_leaves_absent_target_absent(
    tmp_path: Path, wordgen: ModuleType, monkeypatch: pytest.MonkeyPatch
) -> None:
    target = tmp_path / "absent.docx"
    document = Document()
    document.add_paragraph("replacement")

    monkeypatch.setattr(
        wordgen,
        "_validate_docx",
        lambda _path: (_ for _ in ()).throw(ValueError("forced validation failure")),
    )
    with pytest.raises(ValueError, match="forced validation failure"):
        wordgen._save_docx_atomically(document, target)

    assert not target.exists()
    assert list(tmp_path.glob(".*.tmp.docx")) == []


def test_atomic_success_replaces_target_and_removes_staging_file(
    tmp_path: Path, wordgen: ModuleType
) -> None:
    target = tmp_path / "replace.docx"
    target.write_bytes(b"old")
    document = Document()
    document.add_paragraph("new content")

    published, validation = wordgen._save_docx_atomically(document, target)

    assert published == target
    assert validation["package_valid"] is True
    assert Document(str(target)).paragraphs[0].text == "new content"
    assert list(tmp_path.glob(".*.tmp.docx")) == []


def test_subprocess_error_contains_stage_exception_and_traceback(tmp_path: Path) -> None:
    completed = subprocess.run(
        [sys.executable, str(GENERATOR)],
        input=json.dumps(
            {
                "markdown": "# Bad",
                "theme": "not-a-theme",
                "output_path": str(tmp_path / "bad.docx"),
            }
        ),
        text=True,
        capture_output=True,
        check=False,
    )

    assert completed.returncode != 0
    assert "options:" in completed.stderr
    assert "DocumentGenerationError" in completed.stderr
    assert "Traceback" in completed.stderr
    assert not (tmp_path / "bad.docx").exists()
