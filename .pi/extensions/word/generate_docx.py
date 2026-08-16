#!/usr/bin/env python3
"""Generate a professionally styled Word (.docx) document from Markdown.

The TypeScript Word extension sends a JSON specification on stdin. This module
parses Markdown, renders a style-driven Word document, validates the resulting
OOXML package, atomically publishes it, and prints a JSON result on stdout.
"""

from __future__ import annotations

import json
import os
import re
import sys
import tempfile
import traceback
import xml.etree.ElementTree as ET
import zipfile
from collections.abc import Callable, Iterator
from contextlib import contextmanager
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token

# ============================================================
# Theme, color, and typography systems
# ============================================================


@dataclass(frozen=True)
class Theme:
    accent: str
    text_dark: str
    text_muted: str
    heading_font: str
    body_font: str
    mono_font: str


@dataclass(frozen=True)
class ColorTokens:
    text: str
    text_muted: str
    heading: str
    accent: str
    accent_soft: str
    on_accent: str
    link: str
    link_on_accent: str
    background: str
    surface: str
    surface_alt: str
    border: str
    code_background: str


@dataclass(frozen=True)
class TypeScale:
    title: float
    cover_title: float
    subtitle: float
    h1: float
    h2: float
    h3: float
    h4: float
    h5: float
    h6: float
    body: float
    table: float
    caption: float
    code: float
    header_footer: float


THEMES: dict[str, Theme] = {
    "executive": Theme("1F3A5F", "1F2937", "6B7280", "Calibri Light", "Calibri", "Consolas"),
    "modern": Theme("4F46E5", "111827", "6B7280", "Segoe UI", "Segoe UI", "Consolas"),
    "minimal": Theme("111827", "111827", "6B7280", "Arial", "Arial", "Consolas"),
    "editorial": Theme("7C2D12", "1F2937", "6B7280", "Georgia", "Georgia", "Consolas"),
    "tech": Theme("0F766E", "111827", "6B7280", "Segoe UI", "Calibri", "Consolas"),
}

PAGE_SIZES: dict[str, tuple[float, float]] = {
    "letter": (8.5, 11.0),
    "a4": (8.27, 11.69),
}

WHITE = "FFFFFF"
BLACK = "000000"
_HEX_RE = re.compile(r"^[0-9A-Fa-f]{6}$")
_BR_TAG_RE = re.compile(r"<br\s*/?>", re.IGNORECASE)
_REQUIRED_DOCX_PARTS = {
    "[Content_Types].xml",
    "_rels/.rels",
    "word/document.xml",
}

PENNY_BODY = "Penny Body"
PENNY_QUOTE = "Penny Quote"
PENNY_TABLE_HEADER = "Penny Table Header"
PENNY_TABLE_HEADER_ACCENT = "Penny Table Header On Accent"
PENNY_TABLE_BODY = "Penny Table Body"
PENNY_CAPTION = "Penny Caption"
PENNY_CODE_BLOCK = "Penny Code Block"
PENNY_LIST_CONTINUE = "Penny List Continue"
PENNY_DOCUMENT_TITLE = "Penny Document Title"
PENNY_COVER_TITLE = "Penny Cover Title"
PENNY_SUBTITLE = "Penny Subtitle"
PENNY_METADATA = "Penny Metadata"
PENNY_HEADER_FOOTER = "Penny Header Footer"
PENNY_HYPERLINK = "Penny Hyperlink"
PENNY_HYPERLINK_ACCENT = "Penny Hyperlink On Accent"
PENNY_INLINE_CODE = "Penny Inline Code"
PENNY_INLINE_CODE_ACCENT = "Penny Inline Code On Accent"

_BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]


def _hex_tuple(hex_color: str) -> tuple[int, int, int]:
    return tuple(int(hex_color[index : index + 2], 16) for index in (0, 2, 4))  # type: ignore[return-value]


def _tuple_hex(rgb: tuple[float, float, float]) -> str:
    return "".join(f"{max(0, min(255, round(channel))):02X}" for channel in rgb)


def _mix(first: str, second: str, second_amount: float) -> str:
    """Mix two hex colors, where ``second_amount`` is the share of ``second``."""
    first_rgb = _hex_tuple(first)
    second_rgb = _hex_tuple(second)
    mixed = tuple(
        first_channel * (1.0 - second_amount) + second_channel * second_amount
        for first_channel, second_channel in zip(first_rgb, second_rgb, strict=True)
    )
    return _tuple_hex((mixed[0], mixed[1], mixed[2]))


def _relative_luminance(hex_color: str) -> float:
    channels = []
    for value in _hex_tuple(hex_color):
        normalized = value / 255.0
        channels.append(
            normalized / 12.92 if normalized <= 0.04045 else ((normalized + 0.055) / 1.055) ** 2.4
        )
    return 0.2126 * channels[0] + 0.7152 * channels[1] + 0.0722 * channels[2]


def contrast_ratio(foreground: str, background: str) -> float:
    lighter = max(_relative_luminance(foreground), _relative_luminance(background))
    darker = min(_relative_luminance(foreground), _relative_luminance(background))
    return (lighter + 0.05) / (darker + 0.05)


def _ensure_contrast(color: str, background: str, minimum: float = 4.5) -> str:
    if contrast_ratio(color, background) >= minimum:
        return color
    target = (
        BLACK if contrast_ratio(BLACK, background) >= contrast_ratio(WHITE, background) else WHITE
    )
    for step in range(1, 21):
        candidate = _mix(color, target, step / 20.0)
        if contrast_ratio(candidate, background) >= minimum:
            return candidate
    return target


def derive_palette(theme: Theme, accent_override: str | None = None) -> ColorTokens:
    accent = (accent_override or theme.accent).upper()
    on_accent = WHITE if contrast_ratio(WHITE, accent) >= contrast_ratio(BLACK, accent) else BLACK
    text = _ensure_contrast(theme.text_dark, WHITE)
    text_muted = _ensure_contrast(theme.text_muted, WHITE)
    link = _ensure_contrast(accent, WHITE)
    return ColorTokens(
        text=text,
        text_muted=text_muted,
        heading=link,
        accent=accent,
        accent_soft=_mix(accent, WHITE, 0.86),
        on_accent=on_accent,
        link=link,
        link_on_accent=on_accent,
        background=WHITE,
        surface=WHITE,
        surface_alt=_mix(accent, WHITE, 0.95),
        border=_mix(text, WHITE, 0.80),
        code_background=_mix(text, WHITE, 0.95),
    )


def type_scale(body_pt: float) -> TypeScale:
    """Return a monotonic, body-relative scale for all document typography."""
    return TypeScale(
        title=max(24.0, body_pt * 2.2),
        cover_title=max(30.0, body_pt * 2.8),
        subtitle=max(12.0, body_pt * 1.18),
        h1=max(body_pt + 5.0, body_pt * 1.55),
        h2=max(body_pt + 3.0, body_pt * 1.30),
        h3=max(body_pt + 1.5, body_pt * 1.15),
        h4=body_pt,
        h5=max(body_pt - 0.5, 8.0),
        h6=max(body_pt - 1.0, 8.0),
        body=body_pt,
        table=max(body_pt - 0.5, 8.0),
        caption=max(body_pt - 1.0, 8.0),
        code=max(body_pt - 1.5, 8.0),
        header_footer=max(body_pt - 2.0, 8.0),
    )


# ============================================================
# Options
# ============================================================


@dataclass(frozen=True)
class Options:
    title: str | None
    subtitle: str | None
    author: str | None
    date: str | None
    theme_name: str
    theme: Theme
    colors: ColorTokens
    scale: TypeScale
    font_size_pt: float
    line_spacing: float
    line_break_mode: str
    margin_inches: float
    orientation: str
    page_size: str
    title_mode: str
    include_toc: bool
    include_page_numbers: bool
    header_text: str | None
    footer_text: str | None
    table_style: str
    table_layout: str
    output_path: Path
    staging_path: Path | None
    project_root: Path

    @property
    def content_width_in(self) -> float:
        width, height = PAGE_SIZES[self.page_size]
        page_width = height if self.orientation == "landscape" else width
        return page_width - 2 * self.margin_inches

    @property
    def content_height_in(self) -> float:
        width, height = PAGE_SIZES[self.page_size]
        page_height = width if self.orientation == "landscape" else height
        return page_height - 2 * self.margin_inches


def _opt_str(spec: dict[str, Any], key: str) -> str | None:
    value = spec.get(key)
    if value is None or value == "":
        return None
    return str(value)


def _enum(spec: dict[str, Any], key: str, allowed: list[str], default: str) -> str:
    value = str(spec.get(key) or default).lower()
    if value not in allowed:
        raise ValueError(f"{key} must be one of {allowed}, got {value!r}")
    return value


def _number(spec: dict[str, Any], key: str, default: float, low: float, high: float) -> float:
    raw = spec.get(key)
    value = float(default if raw is None else raw)
    if not low <= value <= high:
        raise ValueError(f"{key} must be between {low} and {high}, got {value}")
    return value


def parse_options(spec: dict[str, Any]) -> Options:
    theme_name = _enum(spec, "theme", list(THEMES), "executive")
    theme = THEMES[theme_name]
    accent = _opt_str(spec, "accent_color")
    if accent:
        accent = accent.lstrip("#").upper()
        if not _HEX_RE.fullmatch(accent):
            raise ValueError(f"accent_color must be a 6-digit hex color, got {accent!r}")

    author = _opt_str(spec, "author")
    legacy_cover = bool(spec.get("cover_page", False))
    title_mode = _enum(spec, "title_mode", ["auto", "none", "inline", "cover"], "auto")
    if title_mode == "auto":
        title_mode = "cover" if legacy_cover else "inline"

    date = _opt_str(spec, "date")
    if date is None and (author or title_mode == "cover"):
        date = datetime.now().strftime("%Y-%m-%d")

    output_path = _opt_str(spec, "output_path")
    if not output_path:
        raise ValueError("output_path is required in the generator spec")

    body_size = _number(spec, "font_size_pt", 11.0, 8.0, 14.0)
    return Options(
        title=_opt_str(spec, "title"),
        subtitle=_opt_str(spec, "subtitle"),
        author=author,
        date=date,
        theme_name=theme_name,
        theme=theme,
        colors=derive_palette(theme, accent),
        scale=type_scale(body_size),
        font_size_pt=body_size,
        line_spacing=_number(spec, "line_spacing", 1.15, 1.0, 2.0),
        line_break_mode=_enum(spec, "line_break_mode", ["preserve", "commonmark"], "preserve"),
        margin_inches=_number(spec, "margin_inches", 1.0, 0.4, 2.0),
        orientation=_enum(spec, "orientation", ["portrait", "landscape"], "portrait"),
        page_size=_enum(spec, "page_size", ["letter", "a4"], "letter"),
        title_mode=title_mode,
        include_toc=bool(spec.get("include_toc", False)),
        include_page_numbers=bool(spec.get("include_page_numbers", True)),
        header_text=_opt_str(spec, "header_text"),
        footer_text=_opt_str(spec, "footer_text"),
        table_style=_enum(spec, "table_style", ["banded", "minimal", "grid", "none"], "banded"),
        table_layout=_enum(spec, "table_layout", ["content", "equal"], "content"),
        output_path=Path(output_path).expanduser(),
        staging_path=(
            Path(staging_path).expanduser()
            if (staging_path := _opt_str(spec, "staging_path"))
            else None
        ),
        project_root=Path(_opt_str(spec, "project_root") or Path.cwd()).expanduser().resolve(),
    )


# ============================================================
# Low-level docx helpers
# ============================================================


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def _set_font_slots(rpr: Any, name: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rfonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{slot}"), name)


def _set_run_font(
    run: Any,
    name: str | None = None,
    size_pt: float | None = None,
    color: str | None = None,
) -> None:
    if name:
        run.font.name = name
        _set_font_slots(run._element.get_or_add_rPr(), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = _rgb(color)


def _set_style_font(
    style: Any,
    name: str,
    size_pt: float,
    color: str,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    style.font.name = name
    style.font.size = Pt(size_pt)
    style.font.color.rgb = _rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    _set_font_slots(style._element.get_or_add_rPr(), name)


def _shade_paragraph(paragraph: Any, fill: str) -> None:
    paragraph._p.get_or_add_pPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill}"/>')
    )


def _shade_run_properties(rpr: Any, fill: str) -> None:
    rpr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill}"/>'))


def _paragraph_borders(paragraph: Any, edges: dict[str, tuple[str, int]]) -> None:
    parts = "".join(
        f'<w:{edge} w:val="single" w:sz="{size}" w:space="4" w:color="{color}"/>'
        for edge, (color, size) in edges.items()
    )
    paragraph._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}>{parts}</w:pBdr>'))


def _add_field(paragraph: Any, instruction: str, placeholder: str | None = None) -> None:
    begin = paragraph.add_run()
    begin._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    instr = paragraph.add_run()
    instr._element.append(
        parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {instruction} </w:instrText>')
    )
    if placeholder is not None:
        separator = paragraph.add_run()
        separator._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        )
        placeholder_run = paragraph.add_run(placeholder)
        placeholder_run.italic = True
    end = paragraph.add_run()
    end._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def _add_hyperlink_container(paragraph: Any, url: str) -> Any:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w", "r")} r:id="{relationship_id}"/>')
    paragraph._p.append(hyperlink)
    return hyperlink


def _set_cell_fill(cell: Any, fill: str) -> None:
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill}"/>')
    )


def _set_cell_border_bottom(cell: Any, color: str, size: int) -> None:
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f"</w:tcBorders>"
        )
    )


def _set_table_borders(table: Any, edges: dict[str, tuple[str, int]]) -> None:
    parts = "".join(
        f'<w:{edge} w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        for edge, (color, size) in edges.items()
    )
    table._tbl.tblPr.append(parse_xml(f"<w:tblBorders {nsdecls('w')}>{parts}</w:tblBorders>"))


def _set_table_cell_margins(table: Any) -> None:
    table._tbl.tblPr.append(
        parse_xml(
            f'<w:tblCellMar {nsdecls("w")}>'
            f'<w:top w:w="70" w:type="dxa"/><w:bottom w:w="70" w:type="dxa"/>'
            f'<w:start w:w="100" w:type="dxa"/><w:end w:w="100" w:type="dxa"/>'
            f"</w:tblCellMar>"
        )
    )


def _set_row_policy(row: Any, *, repeat_header: bool) -> None:
    trpr = row._tr.get_or_add_trPr()
    trpr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    if repeat_header:
        trpr.append(parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>'))


def _set_cell_width(cell: Any, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    cell._tc.get_or_add_tcPr().get_or_add_tcW().width = Inches(width_inches)


def _set_page_number_start(section: Any, number: int) -> None:
    section_properties = section._sectPr
    existing = section_properties.find(qn("w:pgNumType"))
    if existing is not None:
        section_properties.remove(existing)
    section_properties.append(parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="{number}"/>'))


def _request_field_update(doc: Any) -> None:
    settings = doc.settings._element
    current = settings.find(qn("w:updateFields"))
    if current is not None:
        settings.remove(current)
    settings.append(parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>'))


def _style(doc: Any, name: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[name]
        return name
    except KeyError:
        return fallback


# ============================================================
# Inline rendering
# ============================================================


@dataclass(frozen=True)
class RunProfile:
    link_style: str = PENNY_HYPERLINK
    code_style: str = PENNY_INLINE_CODE
    italic: bool = False


@dataclass
class _InlineState:
    bold: bool = False
    italic: bool = False
    strike: bool = False
    link: Any = None


_INLINE_TOGGLES: dict[str, tuple[str, bool]] = {
    "strong_open": ("bold", True),
    "strong_close": ("bold", False),
    "em_open": ("italic", True),
    "em_close": ("italic", False),
    "s_open": ("strike", True),
    "s_close": ("strike", False),
}

InlineImageHandler = Callable[[Any, str, str], None]


class InlineRenderer:
    """Render markdown-it inline children into style-inheriting Word runs."""

    def __init__(
        self,
        profile: RunProfile,
        line_break_mode: str,
        image_handler: InlineImageHandler | None = None,
    ) -> None:
        self.profile = profile
        self.line_break_mode = line_break_mode
        self.image_handler = image_handler

    def render(self, paragraph: Any, children: list[Token]) -> None:
        state = _InlineState()
        for token in children:
            self._render_token(paragraph, token, state)

    def _new_run(
        self,
        paragraph: Any,
        state: _InlineState,
        text: str = "",
        *,
        style: str | None = None,
    ) -> Any:
        run = paragraph.add_run(text)
        run.bold = state.bold or None
        run.italic = state.italic or self.profile.italic or None
        run.font.strike = state.strike or None
        if style:
            run.style = style
        elif state.link is not None:
            run.style = self.profile.link_style
        if state.link is not None:
            state.link.append(run._element)
        return run

    def _emit_break(self, paragraph: Any, state: _InlineState) -> None:
        self._new_run(paragraph, state).add_break()

    def _render_token(self, paragraph: Any, token: Token, state: _InlineState) -> None:
        toggle = _INLINE_TOGGLES.get(token.type)
        if toggle is not None:
            setattr(state, toggle[0], toggle[1])
        elif token.type == "text":
            self._new_run(paragraph, state, token.content)
        elif token.type == "code_inline":
            style = self.profile.link_style if state.link is not None else self.profile.code_style
            self._new_run(paragraph, state, token.content, style=style)
        elif token.type == "link_open":
            state.link = _add_hyperlink_container(paragraph, str(token.attrs.get("href", "")))
        elif token.type == "link_close":
            state.link = None
        elif token.type in ("softbreak", "hardbreak", "html_inline"):
            self._render_break_token(paragraph, token, state)
        elif token.type == "image" and self.image_handler is not None:
            alt = token.content or str(token.attrs.get("alt", ""))
            run = self._new_run(paragraph, state)
            self.image_handler(run, str(token.attrs.get("src", "")), alt)

    def _render_break_token(self, paragraph: Any, token: Token, state: _InlineState) -> None:
        if token.type == "softbreak" and self.line_break_mode != "preserve":
            self._new_run(paragraph, state, " ")
        elif token.type != "html_inline" or _BR_TAG_RE.fullmatch(token.content.strip()):
            self._emit_break(paragraph, state)


# ============================================================
# Block rendering
# ============================================================


@dataclass(frozen=True)
class TableCellSpec:
    children: list[Token]
    alignment: str | None


class DocxRenderer:
    """Walk the markdown-it block token stream and emit styled Word content."""

    def __init__(self, doc: Any, opts: Options) -> None:
        self.doc = doc
        self.opts = opts
        self.theme = opts.theme
        self.colors = opts.colors
        self.warnings: list[str] = []
        self.stats = {"headings": 0, "tables": 0, "code_blocks": 0, "images": 0}

    def _inline_renderer(self, profile: RunProfile | None = None) -> InlineRenderer:
        return InlineRenderer(
            profile or RunProfile(),
            self.opts.line_break_mode,
            self._render_inline_image,
        )

    def render(self, tokens: list[Token]) -> None:
        index = 0
        while index < len(tokens):
            index = self._render_block(tokens, index)

    def _render_block(self, tokens: list[Token], index: int) -> int:
        token = tokens[index]
        if token.type == "heading_open":
            return self._render_heading(tokens, index)
        if token.type == "paragraph_open":
            return self._render_paragraph(tokens, index)
        if token.type in ("bullet_list_open", "ordered_list_open"):
            return self._render_list(tokens, index, 0)
        if token.type == "blockquote_open":
            return self._render_blockquote(tokens, index)
        if token.type in ("fence", "code_block"):
            self._render_code(token)
            return index + 1
        if token.type == "hr":
            self._render_hr()
            return index + 1
        if token.type == "table_open":
            return self._render_table(tokens, index)
        return index + 1

    def _render_heading(self, tokens: list[Token], index: int) -> int:
        level = int(tokens[index].tag[1])
        inline = tokens[index + 1]
        paragraph = self.doc.add_paragraph(style=_style(self.doc, f"Heading {min(level, 6)}"))
        self._inline_renderer().render(paragraph, inline.children or [])
        if level == 1:
            _paragraph_borders(paragraph, {"bottom": (self.colors.accent_soft, 6)})
        self.stats["headings"] += 1
        return index + 3

    def _render_paragraph(self, tokens: list[Token], index: int) -> int:
        inline = tokens[index + 1]
        children = inline.children or []
        only_image = len(children) == 1 and children[0].type == "image"
        if only_image:
            self._render_block_image(
                str(children[0].attrs.get("src", "")),
                children[0].content or str(children[0].attrs.get("alt", "")),
            )
            return index + 3
        paragraph = self.doc.add_paragraph(style=PENNY_BODY)
        self._inline_renderer().render(paragraph, children)
        return index + 3

    def _render_list(self, tokens: list[Token], index: int, depth: int) -> int:
        ordered = tokens[index].type == "ordered_list_open"
        close_type = tokens[index].type.replace("_open", "_close")
        number = int(str(tokens[index].attrs.get("start", 1) or 1))
        index += 1
        while tokens[index].type != close_type:
            if tokens[index].type == "list_item_open":
                index = self._render_list_item(tokens, index, depth, ordered, number)
                number += 1
            else:
                index += 1
        return index + 1

    def _render_list_item(
        self,
        tokens: list[Token],
        index: int,
        depth: int,
        ordered: bool,
        number: int,
    ) -> int:
        index += 1
        first_paragraph = True
        while tokens[index].type != "list_item_close":
            if tokens[index].type == "paragraph_open":
                index = self._render_list_paragraph(
                    tokens,
                    index,
                    depth,
                    ordered=ordered,
                    number=number if ordered and first_paragraph else None,
                    continuation=not first_paragraph,
                )
                first_paragraph = False
            elif tokens[index].type in ("bullet_list_open", "ordered_list_open"):
                index = self._render_list(tokens, index, depth + 1)
            else:
                index = self._render_block(tokens, index)
        return index + 1

    def _render_list_paragraph(
        self,
        tokens: list[Token],
        index: int,
        depth: int,
        *,
        ordered: bool,
        number: int | None,
        continuation: bool,
    ) -> int:
        if continuation:
            paragraph = self.doc.add_paragraph(style=PENNY_LIST_CONTINUE)
            paragraph.paragraph_format.left_indent = Inches(0.25 * min(depth, 5) + 0.25)
        elif ordered:
            paragraph = self.doc.add_paragraph(style=PENNY_BODY)
            left = 0.25 * min(depth, 5) + 0.25
            paragraph.paragraph_format.left_indent = Inches(left)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(left), WD_TAB_ALIGNMENT.LEFT)
            paragraph.add_run(f"{number}.\t")
        else:
            paragraph = self.doc.add_paragraph(
                style=_style(self.doc, _BULLET_STYLES[min(depth, len(_BULLET_STYLES) - 1)])
            )
        paragraph.paragraph_format.space_after = Pt(2)
        self._inline_renderer().render(paragraph, tokens[index + 1].children or [])
        return index + 3

    def _render_blockquote(self, tokens: list[Token], index: int) -> int:
        index += 1
        first = True
        while tokens[index].type != "blockquote_close":
            if tokens[index].type == "paragraph_open":
                inline = tokens[index + 1]
                paragraph = self.doc.add_paragraph(style=PENNY_QUOTE)
                paragraph.paragraph_format.space_before = Pt(6 if first else 2)
                _paragraph_borders(paragraph, {"left": (self.colors.accent, 18)})
                self._inline_renderer(RunProfile(italic=True)).render(
                    paragraph, inline.children or []
                )
                first = False
                index += 3
            else:
                index = self._render_block(tokens, index)
        if self.doc.paragraphs:
            self.doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
        return index + 1

    def _render_code(self, token: Token) -> None:
        lines = token.content.rstrip("\n").split("\n")
        paragraph = self.doc.add_paragraph(style=PENNY_CODE_BLOCK)
        for line_index, line in enumerate(lines):
            paragraph.add_run(line)
            if line_index < len(lines) - 1:
                paragraph.add_run().add_break()
        _shade_paragraph(paragraph, self.colors.code_background)
        _paragraph_borders(
            paragraph,
            {
                "top": (self.colors.border, 4),
                "bottom": (self.colors.border, 4),
                "left": (self.colors.border, 4),
                "right": (self.colors.border, 4),
            },
        )
        self.stats["code_blocks"] += 1

    def _render_hr(self) -> None:
        paragraph = self.doc.add_paragraph(style=PENNY_BODY)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(10)
        _paragraph_borders(paragraph, {"bottom": (self.colors.border, 4)})

    def _resolve_image_path(self, source: str) -> Path:
        path = Path(source).expanduser()
        return path if path.is_absolute() else self.opts.project_root / path

    def _image_dimensions_inches(self, path: Path) -> tuple[float, float]:
        max_width = self.opts.content_width_in
        max_height = max(1.0, self.opts.content_height_in - 0.75)
        try:
            from PIL import Image, ImageOps

            with Image.open(path) as image:
                transposed = ImageOps.exif_transpose(image)
                dpi = transposed.info.get("dpi", (96.0, 96.0))
                dpi_x = float(dpi[0] or 96.0)
                dpi_y = float(dpi[1] or 96.0)
                natural_width = transposed.width / dpi_x
                natural_height = transposed.height / dpi_y
            scale = min(1.0, max_width / natural_width, max_height / natural_height)
            return natural_width * scale, natural_height * scale
        except Exception:
            return max_width, min(max_height, max_width * 0.65)

    def _add_picture(self, run: Any, source: str, alt: str) -> bool:
        path = self._resolve_image_path(source)
        if not path.is_file():
            self.warnings.append(f"image not found: {source}")
            run.add_text(f"[image unavailable: {alt or source}]")
            _set_run_font(
                run,
                self.theme.body_font,
                self.opts.scale.caption,
                self.colors.text_muted,
            )
            return False
        width, height = self._image_dimensions_inches(path)
        inline_shape = run.add_picture(str(path), width=Inches(width), height=Inches(height))
        if alt:
            inline_shape._inline.docPr.set("descr", alt)
            inline_shape._inline.docPr.set("title", alt)
        self.stats["images"] += 1
        return True

    def _render_inline_image(self, run: Any, source: str, alt: str) -> None:
        self._add_picture(run, source, alt)

    def _render_block_image(self, source: str, alt: str) -> None:
        paragraph = self.doc.add_paragraph(style=PENNY_BODY)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not self._add_picture(paragraph.add_run(), source, alt):
            return
        if alt:
            paragraph.paragraph_format.keep_with_next = True
            caption = self.doc.add_paragraph(style=PENNY_CAPTION)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.keep_with_next = False
            caption.add_run(alt)

    def _render_table(self, tokens: list[Token], index: int) -> int:
        rows: list[tuple[list[TableCellSpec], bool]] = []
        while tokens[index].type != "table_close":
            if tokens[index].type == "tr_open":
                index, cells, is_header = self._collect_row(tokens, index)
                rows.append((cells, is_header))
            else:
                index += 1
        self._emit_table(rows)
        self.stats["tables"] += 1
        return index + 1

    def _collect_row(
        self, tokens: list[Token], index: int
    ) -> tuple[int, list[TableCellSpec], bool]:
        cells: list[TableCellSpec] = []
        is_header = False
        index += 1
        while tokens[index].type != "tr_close":
            if tokens[index].type in ("th_open", "td_open"):
                cell_token = tokens[index]
                is_header = is_header or cell_token.type == "th_open"
                style = str(cell_token.attrs.get("style", ""))
                alignment_match = re.search(r"text-align\s*:\s*(left|center|right)", style)
                cells.append(
                    TableCellSpec(
                        tokens[index + 1].children or [],
                        alignment_match.group(1) if alignment_match else None,
                    )
                )
                index += 3
            else:
                index += 1
        return index + 1, cells, is_header

    @staticmethod
    def _visible_cell_text(children: list[Token]) -> str:
        parts: list[str] = []
        for token in children:
            if token.type in ("text", "code_inline"):
                parts.append(token.content)
            elif token.type == "image":
                parts.append(token.content or str(token.attrs.get("alt", "")))
            elif token.type in ("softbreak", "hardbreak"):
                parts.append(" ")
        return "".join(parts).strip()

    def _column_widths(
        self, rows: list[tuple[list[TableCellSpec], bool]], columns: int
    ) -> list[float]:
        if self.opts.table_layout == "equal":
            return [self.opts.content_width_in / columns] * columns

        weights: list[float] = []
        for column_index in range(columns):
            lengths = sorted(
                len(self._visible_cell_text(cells[column_index].children))
                for cells, _ in rows
                if column_index < len(cells)
            )
            if not lengths:
                weights.append(6.0)
                continue
            percentile_index = round((len(lengths) - 1) * 0.75)
            weighted_length = 0.7 * lengths[percentile_index] + 0.3 * lengths[-1]
            weights.append(max(6.0, min(50.0, weighted_length)))

        base_width = self.opts.content_width_in * 0.35 / columns
        flexible_width = self.opts.content_width_in - base_width * columns
        total_weight = sum(weights)
        return [base_width + flexible_width * weight / total_weight for weight in weights]

    def _emit_table(self, rows: list[tuple[list[TableCellSpec], bool]]) -> None:
        if not rows:
            return
        columns = max(len(cells) for cells, _ in rows)
        table = self.doc.add_table(rows=len(rows), cols=columns)
        table.autofit = False
        _set_table_cell_margins(table)
        self._apply_table_borders(table)
        widths = self._column_widths(rows, columns)

        for column_index, width in enumerate(widths):
            table.columns[column_index].width = Inches(width)
            table._tbl.tblGrid.gridCol_lst[column_index].w = Inches(width)

        for row_index, (cells, is_header) in enumerate(rows):
            row = table.rows[row_index]
            _set_row_policy(row, repeat_header=is_header)
            band = self.opts.table_style == "banded" and not is_header and row_index % 2 == 0
            for column_index in range(columns):
                cell = row.cells[column_index]
                _set_cell_width(cell, widths[column_index])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                spec = cells[column_index] if column_index < len(cells) else TableCellSpec([], None)
                self._fill_cell(cell, spec, is_header, band)

    def _apply_table_borders(self, table: Any) -> None:
        hairline = (self.colors.border, 4)
        if self.opts.table_style == "banded":
            _set_table_borders(table, {"top": hairline, "bottom": hairline, "insideH": hairline})
        elif self.opts.table_style == "grid":
            _set_table_borders(
                table,
                {
                    edge: hairline
                    for edge in ("top", "bottom", "left", "right", "insideH", "insideV")
                },
            )

    def _fill_cell(self, cell: Any, spec: TableCellSpec, is_header: bool, band: bool) -> None:
        filled_header = is_header and self.opts.table_style in ("banded", "grid")
        if filled_header:
            paragraph_style = PENNY_TABLE_HEADER_ACCENT
            profile = RunProfile(PENNY_HYPERLINK_ACCENT, PENNY_INLINE_CODE_ACCENT)
            _set_cell_fill(cell, self.colors.accent)
        elif is_header:
            paragraph_style = PENNY_TABLE_HEADER
            profile = RunProfile()
            if self.opts.table_style == "minimal":
                _set_cell_border_bottom(cell, self.colors.accent, 12)
        else:
            paragraph_style = PENNY_TABLE_BODY
            profile = RunProfile()
            if band:
                _set_cell_fill(cell, self.colors.surface_alt)

        paragraph = cell.paragraphs[0]
        paragraph.style = paragraph_style
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        alignment = (
            {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(spec.alignment)
            if spec.alignment is not None
            else None
        )
        if alignment is not None:
            paragraph.alignment = alignment
        self._inline_renderer(profile).render(paragraph, spec.children)


# ============================================================
# Document setup
# ============================================================


def _get_or_add_style(doc: Any, name: str, style_type: Any, base: str | None = None) -> Any:
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, style_type)
    if base:
        style.base_style = doc.styles[base]
    style.hidden = False
    return style


def _configure_paragraph_style(
    doc: Any,
    name: str,
    *,
    base: str,
    font: str,
    size: float,
    color: str,
    bold: bool | None = None,
    italic: bool | None = None,
    space_before: float = 0,
    space_after: float = 0,
    line_spacing: float | None = None,
    keep_with_next: bool = False,
) -> Any:
    style = _get_or_add_style(doc, name, WD_STYLE_TYPE.PARAGRAPH, base)
    _set_style_font(style, font, size, color, bold=bold, italic=italic)
    formatting = style.paragraph_format
    formatting.space_before = Pt(space_before)
    formatting.space_after = Pt(space_after)
    formatting.line_spacing = line_spacing
    formatting.keep_with_next = keep_with_next
    return style


def _configure_character_style(
    doc: Any,
    name: str,
    *,
    color: str,
    font: str | None = None,
    size: float | None = None,
    underline: bool = False,
    shading: str | None = None,
) -> Any:
    style = _get_or_add_style(doc, name, WD_STYLE_TYPE.CHARACTER)
    if font is not None:
        style.font.name = font
        _set_font_slots(style._element.get_or_add_rPr(), font)
    if size is not None:
        style.font.size = Pt(size)
    style.font.color.rgb = _rgb(color)
    style.font.underline = underline
    if shading:
        _shade_run_properties(style._element.get_or_add_rPr(), shading)
    return style


def _setup_styles(doc: Any, opts: Options) -> None:
    theme = opts.theme
    colors = opts.colors
    scale = opts.scale

    _set_style_font(doc.styles["Normal"], theme.body_font, scale.body, colors.text)
    normal_format = doc.styles["Normal"].paragraph_format
    normal_format.space_after = Pt(6)
    normal_format.line_spacing = opts.line_spacing

    heading_specs = {
        "Heading 1": (scale.h1, colors.heading, True, False, 16, 6),
        "Heading 2": (scale.h2, colors.text, True, False, 12, 4),
        "Heading 3": (scale.h3, colors.heading, True, False, 10, 3),
        "Heading 4": (scale.h4, colors.text, True, False, 9, 3),
        "Heading 5": (scale.h5, colors.text_muted, True, False, 8, 2),
        "Heading 6": (scale.h6, colors.text_muted, False, True, 7, 2),
    }
    for name, (size, color, bold, italic, before, after) in heading_specs.items():
        style = doc.styles[name]
        _set_style_font(
            style,
            theme.heading_font,
            size,
            color,
            bold=bold,
            italic=italic,
        )
        formatting = style.paragraph_format
        formatting.space_before = Pt(before)
        formatting.space_after = Pt(after)
        formatting.keep_with_next = True

    for name in _BULLET_STYLES:
        _set_style_font(doc.styles[name], theme.body_font, scale.body, colors.text)
        doc.styles[name].paragraph_format.space_after = Pt(2)
        doc.styles[name].paragraph_format.line_spacing = opts.line_spacing

    _configure_paragraph_style(
        doc,
        PENNY_BODY,
        base="Normal",
        font=theme.body_font,
        size=scale.body,
        color=colors.text,
        space_after=6,
        line_spacing=opts.line_spacing,
    )
    quote = _configure_paragraph_style(
        doc,
        PENNY_QUOTE,
        base=PENNY_BODY,
        font=theme.body_font,
        size=scale.body,
        color=colors.text_muted,
        italic=True,
        space_after=2,
        line_spacing=opts.line_spacing,
    )
    quote.paragraph_format.left_indent = Inches(0.25)
    _configure_paragraph_style(
        doc,
        PENNY_TABLE_HEADER,
        base=PENNY_BODY,
        font=theme.body_font,
        size=scale.table,
        color=colors.heading,
        bold=True,
    )
    _configure_paragraph_style(
        doc,
        PENNY_TABLE_HEADER_ACCENT,
        base=PENNY_BODY,
        font=theme.body_font,
        size=scale.table,
        color=colors.on_accent,
        bold=True,
    )
    _configure_paragraph_style(
        doc,
        PENNY_TABLE_BODY,
        base=PENNY_BODY,
        font=theme.body_font,
        size=scale.table,
        color=colors.text,
    )
    _configure_paragraph_style(
        doc,
        PENNY_CAPTION,
        base=PENNY_BODY,
        font=theme.body_font,
        size=scale.caption,
        color=colors.text_muted,
        italic=True,
        space_after=8,
        keep_with_next=False,
    )
    code_style = _configure_paragraph_style(
        doc,
        PENNY_CODE_BLOCK,
        base="Normal",
        font=theme.mono_font,
        size=scale.code,
        color=colors.text,
        space_before=6,
        space_after=6,
        line_spacing=1.0,
    )
    code_style.paragraph_format.left_indent = Inches(0.15)
    code_style.paragraph_format.right_indent = Inches(0.15)
    _configure_paragraph_style(
        doc,
        PENNY_LIST_CONTINUE,
        base=PENNY_BODY,
        font=theme.body_font,
        size=scale.body,
        color=colors.text,
        space_after=2,
        line_spacing=opts.line_spacing,
    )
    _configure_paragraph_style(
        doc,
        PENNY_DOCUMENT_TITLE,
        base="Normal",
        font=theme.heading_font,
        size=scale.title,
        color=colors.heading,
        bold=True,
        space_after=2,
        keep_with_next=True,
    )
    _configure_paragraph_style(
        doc,
        PENNY_COVER_TITLE,
        base="Normal",
        font=theme.heading_font,
        size=scale.cover_title,
        color=colors.heading,
        bold=True,
        space_after=6,
        keep_with_next=True,
    )
    _configure_paragraph_style(
        doc,
        PENNY_SUBTITLE,
        base="Normal",
        font=theme.body_font,
        size=scale.subtitle,
        color=colors.text_muted,
        space_after=2,
        keep_with_next=True,
    )
    _configure_paragraph_style(
        doc,
        PENNY_METADATA,
        base="Normal",
        font=theme.body_font,
        size=scale.caption,
        color=colors.text_muted,
        space_after=2,
    )
    _configure_paragraph_style(
        doc,
        PENNY_HEADER_FOOTER,
        base="Normal",
        font=theme.body_font,
        size=scale.header_footer,
        color=colors.text_muted,
    )

    _configure_character_style(
        doc,
        PENNY_HYPERLINK,
        color=colors.link,
        underline=True,
    )
    _configure_character_style(
        doc,
        PENNY_HYPERLINK_ACCENT,
        color=colors.link_on_accent,
        underline=True,
    )
    _configure_character_style(
        doc,
        PENNY_INLINE_CODE,
        font=theme.mono_font,
        color=colors.text,
        shading=colors.code_background,
    )
    _configure_character_style(
        doc,
        PENNY_INLINE_CODE_ACCENT,
        font=theme.mono_font,
        color=colors.on_accent,
    )


def _setup_page(section: Any, opts: Options) -> None:
    width, height = PAGE_SIZES[opts.page_size]
    if opts.orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        width, height = height, width
    section.page_width = Inches(width)
    section.page_height = Inches(height)
    margin = Inches(opts.margin_inches)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def _setup_header_footer(section: Any, opts: Options) -> None:
    if opts.header_text:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.style = PENNY_HEADER_FOOTER
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.add_run(opts.header_text)

    if not opts.footer_text and not opts.include_page_numbers:
        return
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.style = PENNY_HEADER_FOOTER
    if opts.footer_text:
        paragraph.add_run(opts.footer_text)
        if opts.include_page_numbers:
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(opts.content_width_in), WD_TAB_ALIGNMENT.RIGHT
            )
            paragraph.add_run("\t")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if opts.include_page_numbers:
        _add_field(paragraph, "PAGE")


def _meta_line(opts: Options) -> str | None:
    parts = [part for part in (opts.author, opts.date) if part]
    return "  ·  ".join(parts) if parts else None


def _add_title_block(doc: Any, opts: Options, title: str) -> None:
    paragraph = doc.add_paragraph(style=PENNY_DOCUMENT_TITLE)
    paragraph.add_run(title)
    if opts.subtitle:
        doc.add_paragraph(opts.subtitle, style=PENNY_SUBTITLE)
    metadata = _meta_line(opts)
    if metadata:
        doc.add_paragraph(metadata, style=PENNY_METADATA)
    rule = doc.add_paragraph(style=PENNY_BODY)
    rule.paragraph_format.space_after = Pt(12)
    _paragraph_borders(rule, {"bottom": (opts.colors.accent, 8)})


def _add_cover_page(doc: Any, opts: Options, title: str) -> None:
    available_points = opts.content_height_in * 72.0
    spacer_points = max(72.0, min(180.0, available_points * 0.28))
    spacer = doc.add_paragraph(style=PENNY_BODY)
    spacer.paragraph_format.space_before = Pt(spacer_points)
    bar = doc.add_paragraph(style=PENNY_BODY)
    bar.paragraph_format.space_after = Pt(18)
    bar.paragraph_format.right_indent = Inches(max(opts.content_width_in - 1.6, 0))
    _paragraph_borders(bar, {"top": (opts.colors.accent, 32)})
    doc.add_paragraph(title, style=PENNY_COVER_TITLE)
    if opts.subtitle:
        doc.add_paragraph(opts.subtitle, style=PENNY_SUBTITLE)
    metadata = _meta_line(opts)
    if metadata:
        paragraph = doc.add_paragraph(metadata, style=PENNY_METADATA)
        paragraph.paragraph_format.space_before = Pt(30)


def _add_toc(doc: Any) -> None:
    heading = doc.add_paragraph(style=_style(doc, "Heading 1"))
    heading.add_run("Contents")
    paragraph = doc.add_paragraph(style=PENNY_METADATA)
    _add_field(
        paragraph,
        'TOC \\o "1-3" \\h \\z \\u',
        "[Update the table of contents field if your viewer does not refresh it automatically]",
    )
    _request_field_update(doc)
    doc.add_page_break()


# ============================================================
# Generation, validation, and publication
# ============================================================


class DocumentGenerationError(RuntimeError):
    """A stage-aware document generation failure."""


@contextmanager
def _generation_stage(name: str) -> Iterator[None]:
    try:
        yield
    except DocumentGenerationError:
        raise
    except Exception as exc:
        raise DocumentGenerationError(f"{name}: {type(exc).__name__}: {exc}") from exc


def _load_markdown(spec: dict[str, Any]) -> str:
    markdown = spec.get("markdown")
    markdown_path = spec.get("markdown_path")
    if markdown and str(markdown).strip():
        return str(markdown)
    if markdown_path:
        return Path(str(markdown_path)).read_text(encoding="utf-8")
    raise ValueError("spec requires non-empty 'markdown' or 'markdown_path'")


def _plain_text(inline: Token) -> str:
    return "".join(
        child.content for child in (inline.children or []) if child.type in ("text", "code_inline")
    )


def _derive_title(tokens: list[Token], opts: Options) -> tuple[str, list[Token]]:
    if tokens and tokens[0].type == "heading_open" and tokens[0].tag == "h1":
        heading_text = _plain_text(tokens[1])
        title = opts.title or heading_text or "Document"
        if opts.title_mode != "none" and (opts.title is None or opts.title == heading_text):
            return title, tokens[3:]
        return title, tokens
    return opts.title or "Document", tokens


def _validate_docx(path: Path) -> dict[str, Any]:
    try:
        with zipfile.ZipFile(path, "r") as archive:
            names = set(archive.namelist())
            missing = _REQUIRED_DOCX_PARTS - names
            if missing:
                raise ValueError(f"DOCX package is missing required parts: {sorted(missing)}")
            bad_member = archive.testzip()
            if bad_member is not None:
                raise ValueError(f"DOCX CRC check failed for {bad_member!r}")
            xml_parts = [name for name in names if name.endswith(".xml") or name.endswith(".rels")]
            for name in xml_parts:
                ET.fromstring(archive.read(name))
        Document(str(path))
    except (zipfile.BadZipFile, ET.ParseError) as exc:
        raise ValueError(f"invalid generated DOCX package: {exc}") from exc
    return {
        "package_valid": True,
        "reopen_valid": True,
        "required_parts": sorted(_REQUIRED_DOCX_PARTS),
        "xml_parts_checked": len(xml_parts),
    }


def _save_docx_atomically(
    doc: Any, output_path: Path, staging_path: Path | None = None
) -> tuple[Path, dict[str, Any]]:
    target = output_path.expanduser().resolve()
    target.parent.mkdir(parents=True, exist_ok=True)
    if staging_path is None:
        descriptor, temporary_name = tempfile.mkstemp(
            prefix=f".{target.stem}.", suffix=".tmp.docx", dir=target.parent
        )
        os.close(descriptor)
        temporary_path = Path(temporary_name)
    else:
        temporary_path = staging_path.expanduser().resolve()
        if temporary_path.parent != target.parent or temporary_path == target:
            raise ValueError("staging_path must be a distinct file beside output_path")
        if not temporary_path.exists():
            descriptor = os.open(temporary_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY, 0o600)
            os.close(descriptor)
    try:
        doc.save(str(temporary_path))
        validation = _validate_docx(temporary_path)
        temporary_path.replace(target)
        return target, validation
    finally:
        temporary_path.unlink(missing_ok=True)


def generate(spec: dict[str, Any]) -> dict[str, Any]:
    with _generation_stage("options"):
        opts = parse_options(spec)
    with _generation_stage("load_markdown"):
        markdown = _load_markdown(spec)
    with _generation_stage("parse_markdown"):
        parser = MarkdownIt("commonmark").enable(["table", "strikethrough"])
        tokens = parser.parse(markdown)
        title, body_tokens = _derive_title(tokens, opts)

    with _generation_stage("render"):
        doc = Document()
        doc.core_properties.title = title
        if opts.author:
            doc.core_properties.author = opts.author
        _setup_styles(doc, opts)
        first_section = doc.sections[0]
        _setup_page(first_section, opts)

        if opts.title_mode == "cover":
            _add_cover_page(doc, opts, title)
            body_section = doc.add_section(WD_SECTION.NEW_PAGE)
            _setup_page(body_section, opts)
            _set_page_number_start(body_section, 1)
            _setup_header_footer(body_section, opts)
        else:
            _setup_header_footer(first_section, opts)
            if opts.title_mode == "inline":
                _add_title_block(doc, opts, title)

        if opts.include_toc:
            _add_toc(doc)

        renderer = DocxRenderer(doc, opts)
        renderer.render(body_tokens)

    with _generation_stage("save_validate_publish"):
        output_path, validation = _save_docx_atomically(doc, opts.output_path, opts.staging_path)

    warnings = list(renderer.warnings)
    if opts.include_toc:
        warnings.append(
            "The document contains a Word TOC field; viewers that do not refresh fields may require a manual update."
        )

    return {
        "path": str(output_path),
        "title": title,
        "theme": opts.theme_name,
        "words": len(markdown.split()),
        **renderer.stats,
        "warnings": warnings,
        "validation": validation,
        "normalization": {
            "line_break_mode": opts.line_break_mode,
            "title_mode": opts.title_mode,
            "table_layout": opts.table_layout,
            "leading_h1_consumed": body_tokens is not tokens,
        },
        "resolved_palette": asdict(opts.colors),
        "toc_field_update_requested": opts.include_toc,
    }


def main() -> None:
    spec = json.loads(sys.stdin.read())
    result = generate(spec)
    print(json.dumps(result))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # noqa: BLE001 — one stderr contract with the TS caller
        print(f"word generator failed [{type(exc).__name__}]: {exc}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)
